"""Optional research-note appendices for narrative and audit exports."""

from __future__ import annotations

import csv
import io
import zipfile
from collections.abc import Sequence

from docx import Document

from versevad.exports.docx_report import _normalize_docx
from versevad.research_library import ResearchNote


_FIXED_TIMESTAMP = (1980, 1, 1, 0, 0, 0)


def research_notes_csv(
    notes: Sequence[ResearchNote],
    *,
    include_metadata: bool,
) -> bytes:
    """Return a spreadsheet-friendly notes table."""

    output = io.StringIO(newline="")
    fields = [
        "title",
        "body",
        "anchor_type",
        "anchor_label",
        "module",
        "metric",
        "tags",
        "include_in_export",
    ]
    if include_metadata:
        fields.extend(
            [
                "note_id",
                "parent_type",
                "parent_id",
                "analysis_id",
                "project_id",
                "created_at",
                "updated_at",
            ]
        )
    writer = csv.DictWriter(output, fieldnames=fields)
    writer.writeheader()
    for note in notes:
        row = {
            "title": note.title,
            "body": note.body,
            "anchor_type": note.anchor_type,
            "anchor_label": note.anchor_label,
            "module": note.module,
            "metric": note.metric,
            "tags": " | ".join(note.tags),
            "include_in_export": str(note.include_in_export).lower(),
        }
        if include_metadata:
            row.update(
                {
                    "note_id": note.note_id,
                    "parent_type": note.parent_type,
                    "parent_id": note.parent_id,
                    "analysis_id": note.analysis_id,
                    "project_id": note.project_id,
                    "created_at": note.created_at,
                    "updated_at": note.updated_at,
                }
            )
        writer.writerow(row)
    return b"\xef\xbb\xbf" + output.getvalue().encode("utf-8")


def research_notes_markdown(
    notes: Sequence[ResearchNote],
    *,
    include_metadata: bool,
) -> bytes:
    """Return a readable notes appendix for reproducibility bundles."""

    lines = ["# Research Notes", ""]
    for note in notes:
        lines.extend([f"## {note.title}", ""])
        if note.anchor_label:
            lines.extend([f"Attached to: {note.anchor_label}", ""])
        lines.extend([note.body, ""])
        if note.tags:
            lines.extend([f"Tags: {', '.join(note.tags)}", ""])
        if include_metadata:
            lines.extend(
                [
                    f"Note ID: {note.note_id}",
                    f"Context: {note.parent_type} / {note.parent_id}",
                    f"Created: {note.created_at}",
                    f"Updated: {note.updated_at}",
                    "",
                ]
            )
    return ("\n".join(lines).rstrip() + "\n").encode("utf-8")


def append_research_notes_to_docx(
    report: bytes,
    notes: Sequence[ResearchNote],
    *,
    include_metadata: bool,
) -> bytes:
    """Append selected notes while preserving the report's established design."""

    if not notes:
        return report
    document = Document(io.BytesIO(report))
    document.add_page_break()
    document.add_heading("Research Notes", level=1)
    document.add_paragraph(
        "User-authored interpretive notes selected for this export. They are "
        "separate from VerseVAD's calculated evidence."
    )
    for note in notes:
        document.add_heading(note.title, level=2)
        if note.anchor_label:
            anchor = document.add_paragraph()
            anchor.add_run("Attached to: ").bold = True
            anchor.add_run(note.anchor_label)
        document.add_paragraph(note.body)
        if note.tags:
            tags = document.add_paragraph()
            tags.add_run("Tags: ").bold = True
            tags.add_run(", ".join(note.tags))
        if include_metadata:
            metadata = document.add_paragraph()
            metadata.add_run("Note metadata: ").bold = True
            metadata.add_run(
                f"{note.note_id}; created {note.created_at}; "
                f"updated {note.updated_at}"
            )
    output = io.BytesIO()
    document.save(output)
    return _normalize_docx(output.getvalue())


def add_research_notes_to_audit_bundle(
    bundle: bytes,
    notes: Sequence[ResearchNote],
    *,
    include_metadata: bool,
) -> bytes:
    """Add selected notes and a note-inclusive report to an audit ZIP."""

    if not notes:
        return bundle
    source = io.BytesIO(bundle)
    target = io.BytesIO()
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(
        target,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as output:
        for name in sorted(archive.namelist()):
            content = archive.read(name)
            if name == "VerseVAD_analysis_report.docx":
                content = append_research_notes_to_docx(
                    content,
                    notes,
                    include_metadata=include_metadata,
                )
            information = zipfile.ZipInfo(name, _FIXED_TIMESTAMP)
            information.compress_type = zipfile.ZIP_DEFLATED
            information.external_attr = archive.getinfo(name).external_attr
            output.writestr(information, content)
        for name, content in (
            (
                "research_notes.csv",
                research_notes_csv(notes, include_metadata=include_metadata),
            ),
            (
                "research_notes.md",
                research_notes_markdown(notes, include_metadata=include_metadata),
            ),
        ):
            information = zipfile.ZipInfo(name, _FIXED_TIMESTAMP)
            information.compress_type = zipfile.ZIP_DEFLATED
            output.writestr(information, content)
    return target.getvalue()


__all__ = [
    "add_research_notes_to_audit_bundle",
    "append_research_notes_to_docx",
    "research_notes_csv",
    "research_notes_markdown",
]
