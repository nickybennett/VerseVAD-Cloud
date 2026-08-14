"""Versioned, cross-workspace export contract for VerseVAD.

This module is intentionally a presentation and serialization layer.  It reads
already-computed exporter rows, normalizes their vocabulary, and builds the
shared progressive-disclosure archive layout.  It never computes analytical
measurements.
"""

from __future__ import annotations

import csv
import hashlib
import io
import re
from collections import defaultdict
from collections.abc import Iterable, Mapping, Sequence
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from versevad.exports.metric_labels import human_metric_label


EXPORT_SCHEMA_VERSION = "3.0"
MASTER_METRICS_PATH = "03_MASTER_DATA/Master_Metrics.csv"
EXPORT_METADATA_PATH = "05_REPRODUCIBILITY/Export_Metadata.csv"

MASTER_FIELDS = (
    "export_schema_version",
    "analysis_id",
    "analysis_mode",
    "work_id",
    "title",
    "author",
    "collection",
    "date_label",
    "genre",
    "module_id",
    "metric_id",
    "metric_label",
    "legacy_metric_id",
    "dimension",
    "category",
    "resource_id",
    "resource_label",
    "resource_version",
    "lexical_scope",
    "weighting",
    "analysis_level",
    "corpus_aggregation",
    "value",
    "unit",
    "denominator",
    "eligible_token_count",
    "matched_token_count",
    "unmatched_token_count",
    "token_coverage",
    "eligible_type_count",
    "matched_type_count",
    "unmatched_type_count",
    "type_coverage",
    "observation_count",
    "notes",
)

_SCOPE_ALIASES = {
    "all_lexical": "all_lexical",
    "all lexical": "all_lexical",
    "all lexical tokens": "all_lexical",
    "all_matched": "all_lexical",
    "all matched": "all_lexical",
    "all": "all_lexical",
    "stopword_excluded": "stopword_excluded",
    "stopwords_excluded": "stopword_excluded",
    "stopword excluded": "stopword_excluded",
    "stopwords excluded": "stopword_excluded",
    "content_words": "content_words",
    "content words": "content_words",
    "content words only": "content_words",
}
_WEIGHTING_ALIASES = {
    "token": "token",
    "token-weighted": "token",
    "token weighted": "token",
    "type": "type",
    "type-weighted": "type",
    "type weighted": "type",
}
_ANALYSIS_LEVEL_ALIASES = {
    "": "work",
    "document": "work",
    "poem": "work",
    "work": "work",
    "line": "line",
    "stanza": "stanza",
    "sentence": "sentence",
    "pos": "part_of_speech",
    "part of speech": "part_of_speech",
    "part_of_speech": "part_of_speech",
    "candidate": "candidate",
    "neighbor": "neighbor",
    "corpus": "corpus",
    "collection": "corpus",
}

RESOURCE_REGISTRY: dict[str, dict[str, str]] = {
    "nrc_vad_v2_1": {"label": "NRC VAD Lexicon v2.1", "version": "2.1"},
    "nrc_vad_v1": {"label": "NRC VAD Lexicon v1", "version": "1"},
    "warriner_vad_2013": {"label": "Warriner et al. VAD ratings", "version": "2013"},
    "nrc_emotion_v0_92": {"label": "NRC Emotion Lexicon", "version": "0.92"},
    "nrc_emotion_intensity_v1": {"label": "NRC Emotion Intensity Lexicon", "version": "1"},
    "brysbaert-concreteness-2014": {"label": "Brysbaert concreteness ratings", "version": "2014"},
    "lancaster-sensorimotor-2020": {"label": "Lancaster Sensorimotor Norms", "version": "2020"},
    "subtlex-us-zipf-official": {"label": "SUBTLEX-US Zipf frequencies", "version": "official"},
    "kuperman-aoa-2012-erratum-supplement": {"label": "Kuperman Age of Acquisition ratings", "version": "2012 erratum"},
    "cmudict": {"label": "CMU Pronouncing Dictionary", "version": "bundled"},
}

_RESOURCE_LABEL_ALIASES = {
    "nrc vad lexicon v2.1": "nrc_vad_v2_1",
    "nrc vad lexicon v1": "nrc_vad_v1",
    "warriner et al. vad ratings": "warriner_vad_2013",
    "nrc emotion lexicon": "nrc_emotion_v0_92",
    "nrc emotion intensity lexicon": "nrc_emotion_intensity_v1",
    "brysbaert, warriner, and kuperman concreteness ratings": "brysbaert-concreteness-2014",
    "brysbaert concreteness ratings": "brysbaert-concreteness-2014",
    "lancaster sensorimotor norms": "lancaster-sensorimotor-2020",
    "subtlex-us word frequencies with zipf values": "subtlex-us-zipf-official",
    "kuperman et al. age of acquisition ratings": "kuperman-aoa-2012-erratum-supplement",
}

_MODULE_ALIASES = {
    "lexical_frequency": "frequency",
    "age_of_acquisition": "aoa",
    "sensorimotor_imagery_and_embodiment": "sensorimotor",
    "pronunciation_prosody_foundation": "pronunciation",
    "candidate_meter_and_rhythmic_regularity": "meter",
    "rhyme_and_phonological_patterns": "phonology",
    "vader_sentiment": "vader",
}

_METRIC_EXACT = {
    ("vad", "valence_mean"): "vad.valence.mean",
    ("vad", "arousal_mean"): "vad.arousal.mean",
    ("vad", "dominance_mean"): "vad.dominance.mean",
    ("vad", "vad_mean"): "vad.mean",
    ("concreteness", "concreteness_mean"): "concreteness.mean",
    ("frequency", "frequency_mean"): "frequency.mean_zipf",
    ("aoa", "aoa_mean"): "aoa.mean_years",
    ("word_length", "mean_word_length"): "lexical_style.word_length.mean_characters",
    ("readability", "vv_pre_score"): "readability.vv.pre",
}

_FIXED_SUMMARY_MODULES = {
    "readability_summary.csv": "readability",
    "lexical_style_summary.csv": "lexical_style",
    "pronunciation_summary.csv": "pronunciation",
    "meter_summary.csv": "meter",
    "rhyme_summary.csv": "phonology",
    "vader_sentiment_summary.csv": "vader",
    "poetry_id_summary.csv": "poetry_id",
    "inherited_form_summary.csv": "inherited_form",
    "versemap_summary.csv": "versemap",
    "experiential_dynamics_summary.csv": "experiential_dynamics",
}


def _text(value: object) -> str:
    return "" if value is None else str(value).strip()


def _slug(value: object) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", _text(value).casefold()).strip("_")
    return slug or "unspecified"


def canonical_scope(value: object) -> str:
    text = _text(value).casefold().replace("__", " ")
    return _SCOPE_ALIASES.get(text, _SCOPE_ALIASES.get(text.replace("_", " "), ""))


def canonical_weighting(value: object) -> str:
    text = _text(value).casefold()
    return _WEIGHTING_ALIASES.get(text, _WEIGHTING_ALIASES.get(text.replace("_", " "), ""))


def canonical_analysis_level(value: object) -> str:
    text = _text(value).casefold()
    return _ANALYSIS_LEVEL_ALIASES.get(text, _ANALYSIS_LEVEL_ALIASES.get(text.replace("_", " "), "work"))


def canonical_module_id(value: object) -> str:
    raw = _slug(value)
    return _MODULE_ALIASES.get(raw, raw)


def canonical_resource_id(value: object, label: object = "") -> str:
    raw = _text(value)
    if not raw and not _text(label):
        return ""
    if raw in RESOURCE_REGISTRY:
        return raw
    label_key = _text(label or value).casefold()
    if label_key in _RESOURCE_LABEL_ALIASES:
        return _RESOURCE_LABEL_ALIASES[label_key]
    slug = _slug(raw or label)
    aliases = {
        "brysbaert_concreteness_2014": "brysbaert-concreteness-2014",
        "subtlex_us_zipf_official": "subtlex-us-zipf-official",
        "kuperman_aoa_2012_erratum_supplement": "kuperman-aoa-2012-erratum-supplement",
        "lancaster_sensorimotor_2020": "lancaster-sensorimotor-2020",
    }
    return aliases.get(slug, slug)


def canonical_metric_id(
    module_id: object,
    metric_id: object,
    *,
    dimension: object = "",
    category: object = "",
) -> str:
    module = canonical_module_id(module_id)
    metric = _slug(metric_id)
    dimension_slug = _slug(dimension) if _text(dimension) else ""
    category_slug = _slug(category) if _text(category) else ""

    # Comparison exports historically embedded the resource between module and
    # measurement.  Remove that serialization detail before applying mappings.
    bits = metric.split("_")
    if metric.startswith(module + "_"):
        metric = metric[len(module) + 1 :]
    dotted = _text(metric_id).casefold()
    if "." in dotted:
        dotted_bits = [part for part in dotted.split(".") if part]
        if dotted_bits and canonical_module_id(dotted_bits[0]) == module:
            dotted_bits = dotted_bits[1:]
        resource_slugs = {key.casefold() for key in RESOURCE_REGISTRY}
        dotted_bits = [part for part in dotted_bits if part.casefold() not in resource_slugs]
        metric = "_".join(dotted_bits) or metric

    if module == "vad" and dimension_slug and metric in {"mean", "vad_mean"}:
        return f"vad.{dimension_slug}.mean"
    if module == "vad" and dimension_slug and metric.startswith("vad_"):
        return f"vad.{dimension_slug}.{metric.removeprefix('vad_').replace('_', '.')}"
    if module == "emotion_association":
        name = dimension_slug or category_slug or metric.removesuffix("_association")
        name = name.removesuffix("_association")
        return f"emotion.association.{name}.proportion"
    if module == "emotion_intensity":
        name = dimension_slug or category_slug or metric.removesuffix("_intensity")
        name = name.removesuffix("_intensity")
        return f"emotion.intensity.{name}.mean"
    if module == "sensorimotor":
        name = dimension_slug or category_slug or metric.removeprefix("sensorimotor_").removesuffix("_mean")
        return f"sensorimotor.{name}.mean"
    if module == "concreteness" and "standard_deviation" in metric:
        return "concreteness.population_sd"
    if module == "frequency" and metric in {"mean", "frequency_mean", "mean_zipf"}:
        return "frequency.mean_zipf"
    if module == "aoa" and metric in {"mean", "aoa_mean", "mean_years"}:
        return "aoa.mean_years"
    exact = _METRIC_EXACT.get((module, metric))
    if exact:
        return exact
    return ".".join(part for part in (module, metric.replace("_", ".")) if part)


def _csv_bytes(rows: Sequence[Mapping[str, object]], fields: Sequence[str] | None = None) -> bytes:
    fields = tuple(fields or dict.fromkeys(key for row in rows for key in row))
    if not fields:
        fields = ("record_status",)
    output = io.StringIO(newline="")
    writer = csv.DictWriter(output, fieldnames=list(fields), extrasaction="ignore")
    writer.writeheader()
    writer.writerows(rows)
    return b"\xef\xbb\xbf" + output.getvalue().encode("utf-8")


def _read_csv(content: bytes) -> list[dict[str, str]]:
    if not content:
        return []
    return [dict(row) for row in csv.DictReader(io.StringIO(content.decode("utf-8-sig")))]


def _number_text(value: object) -> str:
    text = _text(value)
    if not text or text.casefold() in {"none", "nan", "not available"}:
        return ""
    return text


def _int_text(value: object) -> str:
    text = _number_text(value)
    if not text:
        return ""
    try:
        return str(int(float(text.replace(",", ""))))
    except ValueError:
        return text


def _coverage_counts(row: Mapping[str, object], weighting: str) -> dict[str, str]:
    eligible_tokens = _int_text(row.get("eligible_token_count", row.get("lexical_tokens", "")))
    matched_tokens = _int_text(row.get("matched_token_count", row.get("matched_tokens", "")))
    unmatched_tokens = _int_text(row.get("unmatched_token_count", ""))
    eligible_types = _int_text(row.get("eligible_type_count", ""))
    matched_types = _int_text(row.get("matched_type_count", ""))
    unmatched_types = _int_text(row.get("unmatched_type_count", ""))
    token_coverage = _number_text(row.get("token_coverage", ""))
    type_coverage = _number_text(row.get("type_coverage", ""))
    # Counts are authoritative. Reconstruct their exact ratio rather than
    # trusting a generic ``coverage`` field, which legacy category metrics
    # sometimes used for the metric proportion itself.
    if eligible_tokens and matched_tokens:
        denominator = int(eligible_tokens)
        token_coverage = str(int(matched_tokens) / denominator) if denominator else ""
    if eligible_types and matched_types:
        denominator = int(eligible_types)
        type_coverage = str(int(matched_types) / denominator) if denominator else ""
    if weighting == "type" and not eligible_types:
        eligible_types, matched_types, type_coverage = eligible_tokens, matched_tokens, token_coverage
        eligible_tokens = matched_tokens = unmatched_tokens = token_coverage = ""
    if eligible_tokens and matched_tokens and not unmatched_tokens:
        unmatched_tokens = str(max(0, int(eligible_tokens) - int(matched_tokens)))
    if eligible_types and matched_types and not unmatched_types:
        unmatched_types = str(max(0, int(eligible_types) - int(matched_types)))
    return {
        "eligible_token_count": eligible_tokens,
        "matched_token_count": matched_tokens,
        "unmatched_token_count": unmatched_tokens,
        "token_coverage": token_coverage,
        "eligible_type_count": eligible_types,
        "matched_type_count": matched_types,
        "unmatched_type_count": unmatched_types,
        "type_coverage": type_coverage,
    }


def _record(
    *,
    analysis_id: str,
    analysis_mode: str,
    work_id: str,
    title: str,
    author: str,
    collection: object = "",
    date_label: object = "",
    genre: object = "",
    module_id: object,
    metric_id: object,
    metric_label: object,
    resource_id: object,
    resource_label: object,
    resource_version: object = "",
    lexical_scope: object = "",
    weighting: object = "",
    analysis_level: object = "work",
    corpus_aggregation: object = "",
    value: object = "",
    unit: object = "",
    denominator: object = "",
    observation_count: object = "",
    notes: object = "",
    coverage_source: Mapping[str, object] | None = None,
    dimension: object = "",
    category: object = "",
    legacy_metric_id: object = "",
) -> dict[str, str]:
    module = canonical_module_id(module_id)
    scope = canonical_scope(lexical_scope)
    weight = canonical_weighting(weighting)
    resource = canonical_resource_id(resource_id, resource_label)
    registered = RESOURCE_REGISTRY.get(resource, {})
    counts = _coverage_counts(coverage_source or {}, weight)
    canonical_id = canonical_metric_id(module, metric_id, dimension=dimension, category=category)
    display_label = human_metric_label(
        module_id=module,
        metric_id=canonical_id,
        fallback=metric_label,
        legacy_metric_id=legacy_metric_id or metric_id,
        dimension=dimension,
        category=category,
    )
    row = {
        "export_schema_version": EXPORT_SCHEMA_VERSION,
        "analysis_id": analysis_id,
        "analysis_mode": analysis_mode,
        "work_id": work_id,
        "title": title,
        "author": author,
        "collection": _text(collection),
        "date_label": _text(date_label),
        "genre": _text(genre),
        "module_id": module,
        "metric_id": canonical_id,
        "metric_label": display_label,
        "legacy_metric_id": _text(legacy_metric_id or metric_id),
        "dimension": _text(dimension),
        "category": _text(category),
        "resource_id": resource,
        "resource_label": _text(resource_label) or registered.get("label", resource),
        "resource_version": _text(resource_version) or registered.get("version", ""),
        "lexical_scope": scope,
        "weighting": weight,
        "analysis_level": canonical_analysis_level(analysis_level),
        "corpus_aggregation": _text(corpus_aggregation),
        "value": _number_text(value),
        "unit": _text(unit),
        "denominator": _text(denominator),
        **counts,
        "observation_count": _int_text(observation_count),
        "notes": _text(notes),
    }
    return {field: row.get(field, "") for field in MASTER_FIELDS}


def _summary_coverage_source(row: Mapping[str, object]) -> dict[str, object]:
    metric = _text(row.get("metric", "")).casefold()
    denominator = _text(row.get("denominator", ""))
    match = re.search(r"(?P<matched>[\d,]+)\s+of\s+(?P<eligible>[\d,]+)", denominator)
    if not match:
        return {}
    matched = match.group("matched").replace(",", "")
    eligible = match.group("eligible").replace(",", "")
    coverage = row.get("value", "") if "coverage" in metric else ""
    if "type" in metric or "unique" in metric:
        return {
            "matched_type_count": matched,
            "eligible_type_count": eligible,
            "type_coverage": coverage,
        }
    return {
        "matched_token_count": matched,
        "eligible_token_count": eligible,
        "token_coverage": coverage,
    }


def _comparison_coverage_source(row: Mapping[str, object]) -> dict[str, object]:
    """Recover comparison counts from the retained denominator description.

    Comparison presentation rows historically serialized only a coverage ratio
    plus a human-readable denominator.  This adapter preserves that already-
    computed evidence while placing token and type counts in their canonical
    fields; it does not rematch or recompute the poem.
    """

    denominator = _text(row.get("denominator", ""))
    match = re.search(r"(?P<matched>[\d,]+)\s+of\s+(?P<eligible>[\d,]+)", denominator)
    if not match:
        return {}
    source: dict[str, object] = {}
    matched = match.group("matched").replace(",", "")
    eligible = match.group("eligible").replace(",", "")
    if canonical_weighting(row.get("weighting", "")) == "type":
        source.update(
            matched_type_count=matched,
            eligible_type_count=eligible,
            type_coverage=row.get("coverage", ""),
        )
    else:
        source.update(
            matched_token_count=matched,
            eligible_token_count=eligible,
            token_coverage=row.get("coverage", ""),
        )
    return source


def _fixed_summary_records(
    files: Mapping[str, bytes],
    *,
    analysis_id: str,
    analysis_mode: str,
    work_id: str,
    title: str,
    author: str,
) -> list[dict[str, str]]:
    """Adapt already-computed fixed-module summary rows into the master schema."""

    records: list[dict[str, str]] = []
    seen_paths: set[str] = set()
    for path, content in files.items():
        basename = Path(path).name.casefold()
        module = _FIXED_SUMMARY_MODULES.get(basename)
        if not module or basename in seen_paths:
            continue
        seen_paths.add(basename)
        for row in _read_csv(content):
            metric = row.get("metric", row.get("metric_id", ""))
            section = _text(row.get("section", "")).casefold()
            if not _text(metric) or section in {"configuration", "method_warning"}:
                continue
            resource_id = "cmudict" if module in {"pronunciation", "meter", "phonology"} else ""
            records.append(
                _record(
                    analysis_id=analysis_id,
                    analysis_mode=analysis_mode,
                    work_id=work_id,
                    title=title,
                    author=author,
                    module_id=module,
                    metric_id=metric,
                    metric_label=_text(metric).replace("_", " ").title(),
                    resource_id=resource_id,
                    resource_label=RESOURCE_REGISTRY.get(resource_id, {}).get("label", ""),
                    value=row.get("value", ""),
                    unit=row.get("unit_or_scale", row.get("unit", "")),
                    denominator=row.get("denominator", ""),
                    notes=row.get("note", ""),
                    coverage_source=_summary_coverage_source(row),
                    legacy_metric_id=metric,
                )
            )
    return records


def canonical_records_from_single(
    files: Mapping[str, bytes],
    *,
    analysis_id: str,
    analysis_mode: str,
    work_id: str,
    title: str,
    author: str,
) -> list[dict[str, str]]:
    candidates = [
        "profile_metrics_all_compatible.csv",
        "05_COMPARATIVE_PROFILES/profile_metrics_all_compatible.csv",
        "profile_metrics_selected.csv",
        "05_COMPARATIVE_PROFILES/profile_metrics_selected.csv",
        "00_START_HERE/profile_comparison.csv",
    ]
    content = next((files[path] for path in candidates if path in files), b"")
    records: list[dict[str, str]] = []
    for row in _read_csv(content):
        records.append(
            _record(
                analysis_id=analysis_id,
                analysis_mode=analysis_mode,
                work_id=work_id,
                title=title,
                author=author,
                module_id=row.get("module_id", ""),
                metric_id=row.get("metric_id", ""),
                metric_label=row.get("metric", row.get("metric_label", "")),
                resource_id=row.get("source_id", ""),
                resource_label=row.get("source", ""),
                lexical_scope=row.get("scope", ""),
                weighting=row.get("weighting", ""),
                value=row.get("value", ""),
                unit=row.get("unit", ""),
                denominator=row.get("denominator", ""),
                observation_count=row.get("observation_count", row.get("observations", "")),
                coverage_source=row,
                legacy_metric_id=row.get("metric_id", ""),
            )
        )
    records.extend(
        _fixed_summary_records(
            files,
            analysis_id=analysis_id,
            analysis_mode=analysis_mode,
            work_id=work_id,
            title=title,
            author=author,
        )
    )
    return _deduplicate(records)


def canonical_records_from_comparison(
    files: Mapping[str, bytes],
    *,
    analysis_id: str,
) -> list[dict[str, str]]:
    records: list[dict[str, str]] = []
    for path, content in files.items():
        if not path.startswith("comparison_") or not path.endswith(".csv"):
            continue
        if path in {"comparison_summary.csv"}:
            continue
        for row in _read_csv(content):
            metric = row.get("metric_id", "")
            module = metric.split(".", 1)[0] if metric else row.get("section", "")
            source = row.get("source", "")
            records.append(
                _record(
                    analysis_id=analysis_id,
                    analysis_mode="compare_poems",
                    work_id=row.get("poem_id", ""),
                    title=row.get("poem_title", ""),
                    author=row.get("author", ""),
                    module_id=module,
                    metric_id=metric,
                    metric_label=row.get("metric", ""),
                    resource_id=source,
                    resource_label=source,
                    lexical_scope=row.get("analysis_view", ""),
                    weighting=row.get("weighting", ""),
                    value=row.get("value", ""),
                    unit=row.get("unit_or_scale", ""),
                    denominator=row.get("denominator", ""),
                    notes=row.get("note", ""),
                    coverage_source=_comparison_coverage_source(row),
                    legacy_metric_id=metric,
                )
            )
    return _deduplicate(records)


def canonical_records_from_corpus(
    files: Mapping[str, bytes],
    *,
    analysis_id: str,
) -> list[dict[str, str]]:
    records: list[dict[str, str]] = []
    for row in _read_csv(files.get("corpus_vad_metrics.csv", b"")):
        module = _module_from_corpus_metric(row.get("metric", ""), row.get("lexicon_id", ""))
        records.append(
            _record(
                analysis_id=analysis_id,
                analysis_mode="corpus",
                work_id=row.get("text_id", ""),
                title=row.get("title", ""),
                author=row.get("author", ""),
                collection=row.get("collection", ""),
                date_label=row.get("date_label", ""),
                genre=row.get("genre", ""),
                module_id=module,
                metric_id=row.get("metric", ""),
                metric_label=" ".join(
                    part for part in (
                        row.get("dimension", "").replace("_", " ").title(),
                        row.get("category", "").replace("_", " ").title(),
                        row.get("metric", "").replace("_", " ").title(),
                    ) if part
                ),
                resource_id=row.get("lexicon_id", ""),
                resource_label=row.get("lexicon", ""),
                lexical_scope=row.get("analysis_view", ""),
                weighting=row.get("weighting", ""),
                value=row.get("value", ""),
                unit=row.get("scale", ""),
                denominator=row.get("denominator", ""),
                observation_count=row.get("observations", ""),
                coverage_source=row,
                dimension=row.get("dimension", ""),
                category=row.get("category", ""),
                legacy_metric_id=row.get("metric", ""),
            )
        )
    for row in _read_csv(files.get("corpus_module_metrics.csv", b"")):
        records.append(
            _record(
                analysis_id=analysis_id,
                analysis_mode="corpus",
                work_id=row.get("text_id", ""),
                title=row.get("title", ""),
                author=row.get("author", ""),
                collection=row.get("collection", ""),
                date_label=row.get("date_label", ""),
                genre=row.get("genre", ""),
                module_id=row.get("module_name", ""),
                metric_id=row.get("metric_id", ""),
                metric_label=row.get("metric_id", "").replace("_", " ").title(),
                resource_id=row.get("module_name", ""),
                resource_label=row.get("module_name", "").replace("_", " ").title(),
                resource_version=row.get("module_version", ""),
                lexical_scope=row.get("scope", ""),
                weighting=row.get("weighting", ""),
                analysis_level=row.get("layer", "work"),
                value=row.get("value", ""),
                unit=row.get("unit", ""),
                denominator=row.get("denominator", ""),
                observation_count=row.get("observation_count", ""),
                notes=row.get("note", ""),
                legacy_metric_id=row.get("metric_id", ""),
            )
        )
    for row in _read_csv(files.get("corpus_module_aggregates.csv", b"")):
        aggregation = _text(row.get("aggregation_method", "")).casefold()
        if "equal" in aggregation or "work" in aggregation:
            aggregation = "equal_work"
        elif "pool" in aggregation or "token" in aggregation or "observation" in aggregation:
            aggregation = "token_pool"
        records.append(
            _record(
                analysis_id=analysis_id,
                analysis_mode="corpus",
                work_id="",
                title="Corpus aggregate",
                author="",
                module_id=row.get("module_name", ""),
                metric_id=row.get("metric_id", ""),
                metric_label=row.get("metric_id", "").replace("_", " ").title(),
                resource_id=row.get("module_name", ""),
                resource_label=row.get("module_name", "").replace("_", " ").title(),
                analysis_level="corpus",
                corpus_aggregation=aggregation,
                value=row.get("value", ""),
                unit=row.get("unit", ""),
                denominator=f"{row.get('works_included', '')} works included",
                observation_count=row.get("observation_count", ""),
                notes=row.get("note", ""),
                legacy_metric_id=row.get("metric_id", ""),
            )
        )
    for row in _read_csv(files.get("corpus_vad_profiles.csv", b"")):
        common = dict(
            analysis_id=analysis_id,
            analysis_mode="corpus",
            work_id="",
            title="Corpus aggregate",
            author="",
            module_id="vad",
            metric_id="vad_mean",
            metric_label=f"{row.get('dimension', '').title()} corpus mean",
            resource_id=row.get("lexicon_id", ""),
            resource_label=row.get("lexicon", ""),
            lexical_scope=row.get("analysis_view", ""),
            weighting=row.get("weighting", ""),
            analysis_level="corpus",
            unit="normalized 0-1",
            dimension=row.get("dimension", ""),
            legacy_metric_id="vad_mean",
        )
        records.append(
            _record(
                **common,
                corpus_aggregation="equal_work",
                value=row.get("work_weighted_volume_mean", ""),
                denominator=f"{row.get('works_included', '')} works included",
                observation_count=row.get("works_included", ""),
                notes="Every eligible work contributes one work-level mean.",
            )
        )
        records.append(
            _record(
                **common,
                corpus_aggregation="token_pool",
                value=row.get("token_weighted_volume_mean", ""),
                denominator=f"{row.get('matched_observations', '')} matched observations pooled",
                observation_count=row.get("matched_observations", ""),
                notes="Eligible matched observations are pooled; longer works can contribute more.",
                coverage_source={
                    "eligible_token_count": row.get("lexical_tokens", ""),
                    "matched_token_count": row.get("matched_observations", ""),
                    "token_coverage": row.get("volume_coverage", ""),
                },
            )
        )
    return _deduplicate(records)


def _module_from_corpus_metric(metric: str, resource_id: str) -> str:
    resource = canonical_resource_id(resource_id)
    if resource in {"nrc_vad_v2_1", "nrc_vad_v1", "warriner_vad_2013"}:
        return "vad"
    if resource == "nrc_emotion_v0_92":
        return "emotion_association"
    if resource == "nrc_emotion_intensity_v1":
        return "emotion_intensity"
    if resource == "brysbaert-concreteness-2014":
        return "concreteness"
    if resource == "subtlex-us-zipf-official":
        return "frequency"
    if resource == "kuperman-aoa-2012-erratum-supplement":
        return "aoa"
    if resource == "lancaster-sensorimotor-2020":
        return "sensorimotor"
    prefix = _slug(metric).split("_", 1)[0]
    return canonical_module_id(prefix)


def _deduplicate(records: Iterable[dict[str, str]]) -> list[dict[str, str]]:
    seen: set[tuple[str, ...]] = set()
    result: list[dict[str, str]] = []
    identity_fields = (
        "analysis_id", "work_id", "module_id", "metric_id", "resource_id",
        "lexical_scope", "weighting", "analysis_level", "corpus_aggregation", "value",
    )
    for row in records:
        key = tuple(row[field] for field in identity_fields)
        if key in seen:
            continue
        seen.add(key)
        result.append(row)
    return result


def _domain(module_id: str) -> str:
    if module_id in {"vad", "emotion_association", "emotion_intensity", "vader", "poetry_id"}:
        return "Affect"
    if module_id in {"concreteness", "sensorimotor"}:
        return "Experience_and_Imagery"
    if module_id in {"frequency", "aoa"}:
        return "Lexical_Accessibility"
    if module_id in {"readability", "lexical_style", "word_length", "structure"}:
        return "Readability_and_Structure"
    if module_id in {"pronunciation", "meter", "phonology", "inherited_form"}:
        return "Sound_and_Form"
    return "Additional_Metrics"


def focused_tables(records: Sequence[Mapping[str, str]], *, mode: str) -> dict[str, bytes]:
    if mode == "compare_poems":
        comparison_groups: dict[str, dict[tuple[str, ...], dict[str, str]]] = defaultdict(dict)
        for row in records:
            domain = _domain(row.get("module_id", ""))
            key = (
                row.get("resource_label", ""),
                row.get("lexical_scope", ""),
                row.get("weighting", ""),
                row.get("metric_label", ""),
                row.get("unit", ""),
            )
            output = comparison_groups[domain].setdefault(
                key,
                {
                    "Resource": key[0],
                    "Scope": key[1],
                    "Weighting": key[2],
                    "Metric": key[3],
                    "Unit": key[4],
                },
            )
            poem_label = row.get("title", "") or row.get("work_id", "")
            output[poem_label] = row.get("value", "")
        files: dict[str, bytes] = {}
        for domain, keyed_rows in comparison_groups.items():
            rows = list(keyed_rows.values())
            for row in rows:
                numeric: list[float] = []
                for key, value in row.items():
                    if key in {"Resource", "Scope", "Weighting", "Metric", "Unit"}:
                        continue
                    try:
                        numeric.append(float(value))
                    except (TypeError, ValueError):
                        pass
                row["Range"] = max(numeric) - min(numeric) if numeric else ""
            files[f"02_METRIC_TABLES/{domain}_Comparison.csv"] = _csv_bytes(rows)
        return files

    groups: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in records:
        coverage = row.get("token_coverage", "") or row.get("type_coverage", "")
        eligible = row.get("eligible_token_count", "") or row.get("eligible_type_count", "")
        matched = row.get("matched_token_count", "") or row.get("matched_type_count", "")
        groups[_domain(row.get("module_id", ""))].append(
            {
                "Work": row.get("title", ""),
                "Author": row.get("author", ""),
                "Resource": row.get("resource_label", ""),
                "Scope": row.get("lexical_scope", ""),
                "Weighting": row.get("weighting", ""),
                "Metric": row.get("metric_label", ""),
                "Value": row.get("value", ""),
                "Unit": row.get("unit", ""),
                "Matched": matched,
                "Eligible": eligible,
                "Coverage": coverage,
                "Aggregation": row.get("corpus_aggregation", ""),
                "Note": row.get("notes", ""),
            }
        )
    suffix = "_Corpus" if mode == "corpus" else ""
    return {
        f"02_METRIC_TABLES/{name}{suffix}.csv": _csv_bytes(rows)
        for name, rows in groups.items()
        if rows
    }


def _coverage_rows(records: Sequence[Mapping[str, str]]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    seen: set[tuple[str, ...]] = set()
    for row in records:
        eligible = row.get("eligible_token_count", "") or row.get("eligible_type_count", "")
        matched = row.get("matched_token_count", "") or row.get("matched_type_count", "")
        unmatched = row.get("unmatched_token_count", "") or row.get("unmatched_type_count", "")
        coverage = row.get("token_coverage", "") or row.get("type_coverage", "")
        if not any((eligible, matched, unmatched, coverage)):
            continue
        key = (
            row.get("work_id", ""), row.get("resource_id", ""), row.get("lexical_scope", ""),
            row.get("weighting", ""), eligible, matched, coverage,
        )
        if key in seen:
            continue
        seen.add(key)
        rows.append(
            {
                "Work": row.get("title", ""),
                "Resource": row.get("resource_label", ""),
                "Scope": row.get("lexical_scope", ""),
                "Weighting": row.get("weighting", ""),
                "Eligible": eligible,
                "Matched": matched,
                "Unmatched": unmatched,
                "Coverage": coverage,
            }
        )
    return rows


def build_coverage_report(
    records: Sequence[Mapping[str, str]],
    *,
    title: str,
    mode_label: str,
) -> bytes:
    detail_rows = _coverage_rows(records)
    grouped: dict[tuple[str, str, str], dict[str, object]] = {}
    for row in detail_rows:
        key = (row["Resource"], row["Scope"], row["Weighting"])
        summary = grouped.setdefault(
            key,
            {
                "Resource": row["Resource"],
                "Scope": row["Scope"],
                "Weighting": row["Weighting"],
                "Works": set(),
                "Eligible": 0,
                "Matched": 0,
                "complete_counts": True,
            },
        )
        if row["Work"]:
            summary["Works"].add(row["Work"])
        try:
            summary["Eligible"] += int(row["Eligible"])
            summary["Matched"] += int(row["Matched"])
        except (TypeError, ValueError):
            summary["complete_counts"] = False
    rows: list[dict[str, str]] = []
    for summary in grouped.values():
        complete = bool(summary.pop("complete_counts"))
        works = summary.pop("Works")
        eligible = int(summary["Eligible"])
        matched = int(summary["Matched"])
        rows.append(
            {
                **{
                    key: str(value)
                    for key, value in summary.items()
                    if key not in {"Eligible", "Matched"}
                },
                "Works": str(len(works)),
                "Eligible": str(eligible) if complete else "",
                "Matched": str(matched) if complete else "",
                "Coverage": (
                    str(matched / eligible) if complete and eligible else ""
                ),
            }
        )
    from versevad.exports.docx_report import _FIXED_CORE_DATE, _normalize_docx

    document = Document()
    document.core_properties.title = "VerseVAD Coverage and Data Quality"
    document.core_properties.subject = mode_label
    document.core_properties.author = "VerseVAD"
    document.core_properties.created = _FIXED_CORE_DATE
    document.core_properties.modified = _FIXED_CORE_DATE
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("VerseVAD\nCoverage and Data Quality")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(13, 19, 43)
    subtitle = document.add_paragraph(f"{title} · {mode_label}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Coverage reports the share of eligible lexical evidence represented by a resource. "
        "Scope-excluded items are not counted as unmatched, and blank values mean unavailable—not zero."
    )
    if not rows:
        document.add_paragraph("No coverage denominator was available for the exported view.")
    else:
        table = document.add_table(rows=1, cols=7)
        table.style = "Light Shading Accent 1"
        headers = ("Resource", "Scope", "Weighting", "Works", "Eligible", "Matched", "Coverage")
        for cell, label in zip(table.rows[0].cells, headers):
            cell.text = label
        for row in rows[:250]:
            cells = table.add_row().cells
            for cell, key in zip(cells, headers):
                value = row.get(key, "")
                if key == "Coverage" and value:
                    try:
                        value = f"{float(value):.1%}"
                    except ValueError:
                        pass
                cell.text = value
        if len(rows) > 250:
            document.add_paragraph(
                f"The readable report shows the first 250 of {len(rows):,} resource/profile summaries. "
                "The complete CSV retains work-level coverage detail."
            )
    document.add_heading("Interpretive safeguards", level=1)
    for text in (
        "Compare token coverage only with token-weighted results and type coverage only with type-weighted results.",
        "Low coverage limits representativeness; it does not imply a neutral lexical score.",
        "Consult Audit for unmatched items and Reproducibility for resource versions and settings.",
    ):
        document.add_paragraph(text, style="List Bullet")
    output = io.BytesIO()
    document.save(output)
    return _normalize_docx(output.getvalue())


def _metric_dictionary(records: Sequence[Mapping[str, str]]) -> bytes:
    rows: list[dict[str, str]] = []
    seen: set[tuple[str, str, str]] = set()
    for row in records:
        key = (row.get("module_id", ""), row.get("metric_id", ""), row.get("resource_id", ""))
        if key in seen:
            continue
        seen.add(key)
        rows.append(
            {
                "module_id": key[0],
                "metric_id": key[1],
                "metric_label": row.get("metric_label", ""),
                "resource_id": key[2],
                "unit": row.get("unit", ""),
                "coverage_fields": "token and/or type counts where applicable",
            }
        )
    return _csv_bytes(rows)


def _resource_manifest(
    records: Sequence[Mapping[str, str]],
    legacy_content: bytes = b"",
) -> bytes:
    legacy = {
        canonical_resource_id(row.get("resource_id", ""), row.get("resource", "")): row
        for row in _read_csv(legacy_content)
    }
    rows: list[dict[str, str]] = []
    seen: set[str] = set()
    for row in records:
        resource_id = row.get("resource_id", "")
        if not resource_id or resource_id in seen:
            continue
        seen.add(resource_id)
        prior = legacy.get(resource_id, {})
        rows.append(
            {
                "resource_id": resource_id,
                "resource": row.get("resource_label", ""),
                "resource_version": prior.get("resource_version", row.get("resource_version", "")),
                "adapter_version": prior.get("adapter_version", ""),
                "source_sha256": prior.get("source_sha256", ""),
                "citation": prior.get(
                    "citation", "See VerseVAD Methodology and the retained module manifest."
                ),
                "license_information": prior.get(
                    "license_information",
                    "See the retained resource manifest and project documentation.",
                ),
            }
        )
    return _csv_bytes(rows)


def _work_summary(records: Sequence[Mapping[str, str]]) -> bytes:
    selected = {
        "vad.valence.mean": "Valence",
        "vad.arousal.mean": "Arousal",
        "vad.dominance.mean": "Dominance",
        "concreteness.mean": "Concreteness",
        "sensorimotor.visual.mean": "Visual Strength",
        "sensorimotor.interoceptive.mean": "Interoceptive Strength",
        "frequency.mean_zipf": "Mean Zipf Frequency",
        "aoa.mean_years": "Mean Age of Acquisition",
        "readability.vv.pre": "Poetry Reading Ease",
    }
    by_work: dict[tuple[str, str, str], dict[str, str]] = {}
    chosen_rank: dict[tuple[tuple[str, str, str], str], int] = {}
    for row in records:
        if row.get("analysis_level") != "work" or not row.get("work_id"):
            continue
        key = (row.get("work_id", ""), row.get("title", ""), row.get("author", ""))
        result = by_work.setdefault(key, {"Work ID": key[0], "Title": key[1], "Author": key[2]})
        label = selected.get(row.get("metric_id", ""))
        if label and row.get("value", ""):
            profile = (row.get("lexical_scope", ""), row.get("weighting", ""))
            rank = {
                ("stopword_excluded", "token"): 0,
                ("content_words", "token"): 1,
                ("all_lexical", "token"): 2,
                ("stopword_excluded", "type"): 3,
                ("content_words", "type"): 4,
                ("all_lexical", "type"): 5,
                ("", ""): 6,
            }.get(profile, 9)
            choice_key = (key, label)
            if rank < chosen_rank.get(choice_key, 99):
                result[label] = row["value"]
                chosen_rank[choice_key] = rank
    return _csv_bytes(list(by_work.values()))


def _corpus_summary(records: Sequence[Mapping[str, str]]) -> bytes:
    """Create an aggregate-first orientation table without recalculating values."""

    grouped: dict[tuple[str, ...], dict[str, str]] = {}
    for row in records:
        if row.get("analysis_level") != "corpus":
            continue
        key = (
            _domain(row.get("module_id", "")),
            row.get("resource_label", ""),
            row.get("lexical_scope", ""),
            row.get("weighting", ""),
            row.get("metric_id", ""),
            row.get("metric_label", ""),
            row.get("unit", ""),
        )
        output = grouped.setdefault(
            key,
            {
                "Profile Area": key[0].replace("_", " "),
                "Resource": key[1],
                "Scope": key[2],
                "Weighting": key[3],
                "Metric": key[5],
                "Equal-work Value": "",
                "Token-pool Value": "",
                "Other Aggregate Value": "",
                "Unit": key[6],
                "Works Represented": "",
                "Pooled Observations": "",
                "Other Aggregate Observations": "",
                "Interpretive Note": "",
            },
        )
        aggregation = row.get("corpus_aggregation", "")
        if aggregation == "equal_work":
            output["Equal-work Value"] = row.get("value", "")
            output["Works Represented"] = row.get("observation_count", "")
        elif aggregation == "token_pool":
            output["Token-pool Value"] = row.get("value", "")
            output["Pooled Observations"] = row.get("observation_count", "")
        else:
            output["Other Aggregate Value"] = row.get("value", "")
            output["Other Aggregate Observations"] = row.get(
                "observation_count", ""
            )
        note = row.get("notes", "")
        if note:
            output["Interpretive Note"] = note
    return _csv_bytes(list(grouped.values()))


def _start_here(*, mode_label: str, export_mode: str, title: str) -> bytes:
    return (
        "VerseVAD Export\n"
        "===============\n\n"
        f"Export schema: {EXPORT_SCHEMA_VERSION}\n"
        f"Analysis mode: {mode_label}\n"
        f"Export type: {'Complete Audit' if export_mode == 'complete_audit' else 'Current View'}\n"
        f"Title: {title or 'Untitled'}\n\n"
        "Start with 01_REPORTS for readable findings and coverage guidance.\n"
        "For a corpus, begin with 02_METRIC_TABLES/Corpus_Summary.csv, then use 03_MASTER_DATA/Work_Summary.csv to orient to individual works.\n"
        "Use the remaining 02_METRIC_TABLES files for detailed, human-readable work-level measurements.\n"
        "Use 03_MASTER_DATA/Master_Metrics.csv as the authoritative machine interface.\n"
        "Use 04_AUDIT to inspect detailed retained evidence.\n"
        "Use 05_REPRODUCIBILITY for settings, resources, metric definitions, warnings, and manifests.\n\n"
        "Current View is an exact projection of the active report section and profile selections.\n"
        "Complete Audit contains all exported profiles and retained evidence. No export layer recalculates metrics.\n"
        "Blank values mean unavailable, not zero.\n"
    ).encode("utf-8")


def _file_guide() -> bytes:
    return _csv_bytes(
        [
            {
                "Folder": "00_READ_ME",
                "Purpose": "Orientation and a concise map of the export.",
                "Start here": "START_HERE.txt",
            },
            {
                "Folder": "01_REPORTS",
                "Purpose": "Readable scholarly reports and coverage guidance.",
                "Start here": "The mode-specific Word report",
            },
            {
                "Folder": "02_METRIC_TABLES",
                "Purpose": "Aggregate-first corpus summary plus detailed human-readable tables grouped by analytical domain.",
                "Start here": "Corpus_Summary.csv, then the domain matching the research question",
            },
            {
                "Folder": "03_MASTER_DATA",
                "Purpose": "Authoritative normalized machine data and curated work summaries.",
                "Start here": "Master_Metrics.csv",
            },
            {
                "Folder": "04_AUDIT",
                "Purpose": "Detailed retained evidence, token records, and module-specific files.",
                "Start here": "Only when inspecting or challenging a result",
            },
            {
                "Folder": "05_REPRODUCIBILITY",
                "Purpose": "Settings, resources, metric definitions, warnings, and file integrity.",
                "Start here": "Export_Metadata.csv and Resource_Manifest.csv",
            },
        ]
    )


def _comparison_summary(records: Sequence[Mapping[str, str]]) -> bytes:
    by_metric: dict[tuple[str, str, str, str, str], list[Mapping[str, str]]] = defaultdict(list)
    for row in records:
        if row.get("analysis_level") != "work":
            continue
        key = (
            row.get("module_id", ""),
            row.get("metric_id", ""),
            row.get("resource_id", ""),
            row.get("lexical_scope", ""),
            row.get("weighting", ""),
        )
        by_metric[key].append(row)
    rows: list[dict[str, str]] = []
    for key, items in by_metric.items():
        numeric_values: list[float] = []
        for item in items:
            try:
                numeric_values.append(float(item.get("value", "")))
            except (TypeError, ValueError):
                continue
        rows.append(
            {
                "Module": key[0],
                "Metric ID": key[1],
                "Resource": items[0].get("resource_label", ""),
                "Scope": key[3],
                "Weighting": key[4],
                "Works represented": len({item.get("work_id", "") for item in items}),
                "Minimum": min(numeric_values) if numeric_values else "",
                "Maximum": max(numeric_values) if numeric_values else "",
                "Range": max(numeric_values) - min(numeric_values) if numeric_values else "",
                "Unit": items[0].get("unit", ""),
            }
        )
    return _csv_bytes(rows)


def standardize_export_files(
    files: Mapping[str, bytes],
    *,
    analysis_mode: str,
    export_mode: str,
    analysis_id: str,
    title: str,
    author: str = "",
    work_id: str = "",
    main_report_path: str,
    main_report_name: str,
) -> dict[str, bytes]:
    """Return the shared schema-v3 archive layout around retained exporter data."""

    source_files = dict(files)
    if analysis_mode in {"single_poem", "other_text"}:
        records = canonical_records_from_single(
            source_files,
            analysis_id=analysis_id,
            analysis_mode=analysis_mode,
            work_id=work_id,
            title=title,
            author=author,
        )
    elif analysis_mode == "compare_poems":
        records = canonical_records_from_comparison(source_files, analysis_id=analysis_id)
    elif analysis_mode == "corpus":
        records = canonical_records_from_corpus(source_files, analysis_id=analysis_id)
    else:
        raise ValueError(f"Unsupported analysis mode: {analysis_mode}")

    report = source_files.pop(main_report_path)
    mode_label = {
        "single_poem": "Single Poem",
        "other_text": "Other Text / Prose",
        "compare_poems": "Compare Poems",
        "corpus": "Corpus / Research Project",
    }[analysis_mode]
    legacy_resource_manifest = next(
        (
            content
            for path, content in source_files.items()
            if Path(path).name.casefold() == "resource_manifest.csv"
        ),
        b"",
    )
    organized: dict[str, bytes] = {
        "00_READ_ME/START_HERE.txt": _start_here(
            mode_label=mode_label, export_mode=export_mode, title=title
        ),
        "00_READ_ME/File_Guide.csv": _file_guide(),
        f"01_REPORTS/{main_report_name}": report,
        "01_REPORTS/Coverage_and_Data_Quality.docx": build_coverage_report(
            records, title=title, mode_label=mode_label
        ),
        MASTER_METRICS_PATH: _csv_bytes(records, MASTER_FIELDS),
        "03_MASTER_DATA/Work_Summary.csv": _work_summary(records),
        "05_REPRODUCIBILITY/Metric_Dictionary.csv": _metric_dictionary(records),
        "05_REPRODUCIBILITY/Resource_Manifest.csv": _resource_manifest(
            records, legacy_resource_manifest
        ),
    }
    organized.update(focused_tables(records, mode=analysis_mode))
    if analysis_mode == "compare_poems":
        organized["03_MASTER_DATA/Comparison_Summary.csv"] = _comparison_summary(records)
    if analysis_mode == "corpus":
        organized["02_METRIC_TABLES/Corpus_Summary.csv"] = _corpus_summary(records)
        aggregate_rows = [
            {
                "Resource": row.get("resource_label", ""),
                "Scope": row.get("lexical_scope", ""),
                "Weighting": row.get("weighting", ""),
                "Metric": row.get("metric_label", ""),
                "Aggregation": row.get("corpus_aggregation", ""),
                "Value": row.get("value", ""),
                "Unit": row.get("unit", ""),
                "Observations": row.get("observation_count", ""),
                "Note": row.get("notes", ""),
            }
            for row in records
            if row.get("analysis_level") == "corpus"
        ]
        organized["03_MASTER_DATA/Corpus_Aggregates.csv"] = _csv_bytes(aggregate_rows)

    repro_names = {
        "reproducibility_readme.txt": "REPRODUCIBILITY_README.txt",
        "resource_manifest.csv": "Legacy_Resource_Manifest.csv",
        "metric_dictionary.csv": "Legacy_Metric_Dictionary.csv",
        "warnings.csv": "Warnings.csv",
        "module_scope_overrides.csv": "Module_Scope_Overrides.csv",
        "corpus_methodology.csv": "Methodology.csv",
        "corpus_project.csv": "Project_Settings.csv",
    }
    master_names = {
        "profile_metrics_all_compatible.csv": "All_Profiles.csv",
        "profile_metrics_selected.csv": "Selected_Profiles.csv",
        "profile_comparison.csv": "All_Profiles.csv",
        "corpus_vad_profiles.csv": "Corpus_Aggregates.csv",
        "corpus_module_aggregates.csv": "Module_Aggregates.csv",
        "corpus_works.csv": "Works.csv",
        "corpus_scope_token_counts.csv": "Scope_Token_Counts.csv",
    }
    for path, content in source_files.items():
        basename = Path(path).name
        lower = basename.casefold()
        if lower in {"file_inventory.txt", "file_inventory.csv", "master_manifest.csv"}:
            continue
        if lower in repro_names:
            target = f"05_REPRODUCIBILITY/{repro_names[lower]}"
        elif lower in master_names:
            target = f"03_MASTER_DATA/{master_names[lower]}"
        else:
            target = f"04_AUDIT/{path}"
        if target not in organized:
            organized[target] = content
        else:
            # A canonical presentation/master file may intentionally replace a
            # legacy summary at the same friendly name. Preserve the original
            # bytes in Audit so the refactor never discards evidence.
            audit_target = f"04_AUDIT/{path}"
            if audit_target not in organized:
                organized[audit_target] = content

    coverage_rows = _coverage_rows(records)
    organized["02_METRIC_TABLES/Coverage_and_Data_Quality.csv"] = _csv_bytes(coverage_rows)
    metadata = [{
        "export_schema_version": EXPORT_SCHEMA_VERSION,
        "analysis_mode": analysis_mode,
        "export_type": export_mode,
        "analysis_id": analysis_id,
        "title": title,
        "generated_at_utc": "",
        "canonical_master_path": MASTER_METRICS_PATH,
        "calculation_policy": "serialization only; no export-time analytical recalculation",
    }]
    organized[EXPORT_METADATA_PATH] = _csv_bytes(metadata)

    # The manifest is built last and is itself intentionally omitted from its
    # checksum list, avoiding recursive content changes.
    manifest_rows = [
        {
            "path": path,
            "export_schema_version": EXPORT_SCHEMA_VERSION,
            "analysis_mode": analysis_mode,
            "export_type": export_mode,
            "bytes": len(content),
            "sha256": hashlib.sha256(content).hexdigest(),
        }
        for path, content in sorted(organized.items())
    ]
    organized["05_REPRODUCIBILITY/Export_Manifest.csv"] = _csv_bytes(manifest_rows)
    organized["05_REPRODUCIBILITY/FILE_INVENTORY.csv"] = _csv_bytes(
        [
            {
                "path": path,
                "folder": path.split("/", 1)[0],
                "bytes": len(content),
                "sha256": hashlib.sha256(content).hexdigest(),
            }
            for path, content in sorted(organized.items())
        ]
    )
    return organized


__all__ = [
    "EXPORT_METADATA_PATH",
    "EXPORT_SCHEMA_VERSION",
    "MASTER_FIELDS",
    "MASTER_METRICS_PATH",
    "RESOURCE_REGISTRY",
    "canonical_analysis_level",
    "canonical_metric_id",
    "canonical_module_id",
    "canonical_resource_id",
    "canonical_scope",
    "canonical_weighting",
    "standardize_export_files",
]
