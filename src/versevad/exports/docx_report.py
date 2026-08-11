"""Shared deterministic narrative Word reports for VerseVAD exports."""

from __future__ import annotations

import csv
import io
import zipfile
from collections import OrderedDict
from dataclasses import dataclass
from datetime import UTC, datetime
from math import isfinite
from typing import Iterable, Mapping, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_FIXED_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
_FIXED_CORE_DATE = datetime(2000, 1, 1, tzinfo=UTC)
_BLUE = "1F4E78"
_DARK_BLUE = "17365D"
_SLATE_BLUE = "466B82"
_PALE_BLUE = "DCE8EF"
_PALE_GRAY = "F2F4F7"
_LIGHTER_BLUE = "EDF3F7"


@dataclass(frozen=True, slots=True)
class ComprehensiveReportFamily:
    """One readable family in the comprehensive analysis report."""

    family_id: str
    title: str
    explanation: str
    cautions: tuple[str, ...] = ()


COMPREHENSIVE_REPORT_FAMILIES: tuple[ComprehensiveReportFamily, ...] = (
    ComprehensiveReportFamily(
        "vad",
        "Affective Lexical Profile: Valence, Arousal, and Dominance",
        "Valence describes normative positivity or pleasantness, arousal describes "
        "activation or intensity, and dominance describes normative power or control. "
        "These are lexical tendencies in the matched evidence, not declarations of "
        "the poem's, speaker's, author's, or reader's emotion.",
        ("The normalized display scale runs from 0 to 1 with a midpoint of 0.5.",),
    ),
    ComprehensiveReportFamily(
        "emotion",
        "Affective Lexical Profile: Emotion Associations, Intensity, and Sentiment",
        "Association proportions record how often eligible evidence is linked with "
        "documented emotion categories. Intensity values summarize the strength of "
        "those lexical associations when the source supplies it.",
        (
            "Association categories can overlap and therefore need not sum to 100%.",
            "VADER is a prose/social-media-oriented rule system and is exploratory for poetry.",
        ),
    ),
    ComprehensiveReportFamily(
        "experience_imagery",
        "Experience and Imagery: Concreteness and Sensorimotor Evidence",
        "Concreteness estimates normative sensory or experience-based grounding. "
        "Lancaster dimensions describe context-free normative associations with "
        "perceptual modalities and bodily action effectors.",
        (
            "Normative lexical affordances support interpretation but do not guarantee imagery or experience for every reader.",
        ),
    ),
    ComprehensiveReportFamily(
        "lexical_accessibility",
        "Lexical Accessibility and Style",
        "SUBTLEX Zipf values describe corpus frequency, with higher values indicating "
        "more common words. Age of Acquisition records retrospective normative age "
        "estimates for matched vocabulary. VerseVAD Poetic Reading Ease estimates "
        "surface-level linguistic accessibility without sentence length. ",
        (
            "Configured common/rare and early/later bands are orientation aids, not universal linguistic boundaries.",
            "Readability scores do not measure thematic, symbolic, interpretive, or literary complexity.",
        ),
    ),
    ComprehensiveReportFamily(
        "readability",
        "Traditional Readability and Processing Demand",
        "VerseVAD Poetic Reading Ease estimates surface-level linguistic accessibility "
        "without sentence length. Traditional readability formulas remain prose-oriented "
        "descriptive evidence and can behave unusually on poetic syntax and lineation.",
        ("Readability scores do not measure thematic, symbolic, interpretive, or literary complexity.",),
    ),
    ComprehensiveReportFamily(
        "structure",
        "Lexical Diversity and Formal Structure",
        "These measures describe vocabulary recurrence, word length, line and stanza "
        "dimensions, and other preserved structural features. Diversity measures respond "
        "differently to text length and token order, so their parameters and scope matter.",
    ),
    ComprehensiveReportFamily(
        "sound_form",
        "Prosody, Rhythm, Sound, and Inherited Form",
        "Pronunciation-supported results describe syllables, stress, candidate meter, "
        "rhyme, and recurring sound evidence. Inherited-form results compare the poem "
        "with versioned rule-based profiles.",
        (
            "Dictionary pronunciations and automatic scansion are analytical candidates, not mandatory performances.",
            "Form confidence and consistency are rule-based indices, not probabilities or declarations of genre identity.",
        ),
    ),
    ComprehensiveReportFamily(
        "poetry_id",
        "PoetryID Lexical-Affective Profile",
        "PoetryID summarizes a versioned lexical-affective profile. Category fit applies "
        "the documented region rules; nearest centroid reports the closest candidate in "
        "the registered profile space.",
        ("The labels are descriptive candidates, not diagnoses, genres, or judgments of quality.",),
    ),
    ComprehensiveReportFamily(
        "versemap",
        "VerseMap Comparative Position",
        "VerseMap positions the poem under Standard Profile 1.0 relative to the selected "
        "reference corpus. Full-space distance determines neighbors; PCA coordinates are "
        "two-dimensional display composites.",
        ("Proximity is descriptive and is not evidence of authorship, influence, quality, or meaning.",),
    ),
)


_PROFILE_FAMILY_BY_MODULE = {
    "vad": "vad",
    "emotion_association": "emotion",
    "emotion_intensity": "emotion",
    "concreteness": "experience_imagery",
    "frequency": "lexical_accessibility",
    "aoa": "lexical_accessibility",
    "sensorimotor": "experience_imagery",
    "word_length": "structure",
}


_CALCULATED_FAMILY_BY_MODULE = {
    **_PROFILE_FAMILY_BY_MODULE,
    "vader_sentiment": "emotion",
    "readability": "readability",
    "lexical_style": "structure",
    "pronunciation": "sound_form",
    "meter": "sound_form",
    "performance_meter": "sound_form",
    "phonology": "sound_form",
    "inherited_form": "sound_form",
    "poetry_id": "poetry_id",
    "versemap": "versemap",
    "lexical_frequency": "lexical_accessibility",
    "age_of_acquisition": "lexical_accessibility",
    "sensorimotor_imagery_and_embodiment": "experience_imagery",
    "pronunciation_prosody_foundation": "sound_form",
    "candidate_meter_and_rhythmic_regularity": "sound_form",
    "rhyme_and_phonological_patterns": "sound_form",
}


_FILE_FAMILY_HINTS: tuple[tuple[tuple[str, ...], str], ...] = (
    (("phase2_vad", "vad_by_part", "lexical_trajectory"), "vad"),
    (("phase2_emotion", "vader_sentiment"), "emotion"),
    (("concreteness", "sensorimotor"), "experience_imagery"),
    (("frequency", "lexical_frequency", "aoa", "age_of_acquisition"), "lexical_accessibility"),
    (("readability",), "readability"),
    (("lexical_style",), "structure"),
    (("pronunciation", "meter", "rhyme", "phonological", "inherited_form"), "sound_form"),
    (("poetry_id",), "poetry_id"),
    (("versemap",), "versemap"),
)


_PROFILE_BACKED_REPORT_FILES = {
    "vad_by_part_of_speech.csv",
    "lexical_trajectory.csv",
    "phase2_emotion_associations.csv",
    "phase2_emotion_intensity.csv",
}


_REPORT_COLUMN_WHITELISTS: Mapping[str, tuple[str, ...]] = {
    "vad_by_part_of_speech.csv": (
        "lexicon",
        "analysis_view",
        "part_of_speech",
        "matched_observations",
        "lexical_token_coverage",
        "token_weighted_mean_valence_0_1",
        "token_weighted_mean_arousal_0_1",
        "token_weighted_mean_dominance_0_1",
        "type_weighted_mean_valence_0_1",
        "type_weighted_mean_arousal_0_1",
        "type_weighted_mean_dominance_0_1",
        "sparse_below_configured_minimum",
    ),
    "phase2_emotion_associations.csv": (
        "category",
        "associated_token_count",
        "associated_unique_type_count",
        "proportion_of_lexical_tokens",
        "proportion_of_matched_emotion_bearing_tokens",
        "proportion_of_unique_lexical_types",
    ),
    "phase2_emotion_intensity.csv": (
        "category",
        "matched_token_occurrences",
        "prevalence_among_lexical_tokens",
        "prevalence_among_emotion_intensity_matches",
        "token_mean",
        "token_median",
        "token_population_standard_deviation",
        "type_mean",
        "type_median",
        "type_population_standard_deviation",
    ),
    "lexical_trajectory.csv": (
        "lexicon_id",
        "lexicon",
        "analysis_view",
        "line_number",
        "stanza_number",
        "source_text",
        "mean_valence_0_1",
        "mean_arousal_0_1",
        "mean_dominance_0_1",
        "mean_concreteness_normalized_0_1",
        "mean_concreteness_source_scale_1_5",
        "vad_matched_observations",
        "concreteness_matched_tokens",
    ),
}


_ATOMIC_FILE_MARKERS = (
    "token_audit",
    "match_audit",
    "observations",
    "processing_tokens",
    "processing_dependencies",
    "processing_entities",
    "processing_orthographic_spans",
    "_lines.csv",
    "_pairs.csv",
    "_internal.csv",
    "alignment_operations",
    "scholar_revisions",
    "_terms.csv",
    "_neighbors.csv",
    "_types.csv",
    "_words.csv",
    "_sentences.csv",
)


@dataclass(frozen=True, slots=True)
class NarrativeReportProfile:
    """Writing guidance for one exported analysis report."""

    title: str
    scope: str
    interpretation: str
    cautions: tuple[str, ...] = ()


REPORT_PROFILES: Mapping[str, NarrativeReportProfile] = {
    "concreteness": NarrativeReportProfile(
        "Concreteness Analysis Report",
        "This report summarizes matched normative concreteness ratings for the "
        "eligible language in the analyzed text.",
        "Higher source ratings indicate language judged as more concrete. Results "
        "describe the words covered by the resource; they do not determine a "
        "text's meaning or literary quality.",
        ("Unmatched items remain missing rather than being assigned a neutral value.",),
    ),
    "frequency": NarrativeReportProfile(
        "Word Frequency and Rarity Report",
        "This report summarizes SUBTLEX-US frequency evidence for eligible words "
        "in the analyzed text.",
        "Higher Zipf values indicate more frequent words in the reference corpus; "
        "lower values indicate rarer matched words.",
        ("Coverage and the configured eligibility scope should accompany comparisons.",),
    ),
    "aoa": NarrativeReportProfile(
        "Age of Acquisition Report",
        "This report summarizes Kuperman et al. normative age-of-acquisition "
        "ratings for matched words in the analyzed text.",
        "Higher values indicate later reported acquisition in the source norms. "
        "They are lexical norms, not estimates of a reader's actual age.",
        ("The source resource is principally a content-word resource.",),
    ),
    "vader_sentiment": NarrativeReportProfile(
        "VADER Sentiment Analysis Report",
        "This report summarizes offline VADER rule-based polarity evidence for "
        "the complete preserved text and its model-segmented sentences.",
        "Positive, neutral, and negative proportions describe VADER's raw lexical "
        "categories; compound is a rule-adjusted normalized composite.",
        (
            "VADER was designed for social-media sentiment and can misread poetic ambiguity, irony, persona, and historical usage.",
            "Its polarity outputs are not declarations of the poem's emotion or a reader's response.",
        ),
    ),
    "readability": NarrativeReportProfile(
        "Readability and Grade-Formula Report",
        "This report summarizes the experimental VerseVAD Poetic Reading Ease "
        "composite and familiar English readability formulas using auditable inputs.",
        "VV-PRE estimates surface-level linguistic accessibility without sentence "
        "length; the traditional scores remain prose-oriented orientation evidence.",
        (
            "VV-PRE does not measure thematic, symbolic, interpretive, or literary complexity.",
            "Its separate High, Moderate, or Limited evidence-confidence label reflects declared coverage and matched-token thresholds; it does not alter the score or represent statistical certainty.",
            "Poetic lineation, fragments, and deliberate syntactic disruption can make the formulas unstable.",
            "Out-of-dictionary syllables use an explicitly labeled orthographic heuristic unless a session override is supplied.",
        ),
    ),
    "sensorimotor imagery and embodiment": NarrativeReportProfile(
        "Sensorimotor Imagery and Embodiment Report",
        "This report summarizes Lancaster Sensorimotor Norms evidence for matched "
        "words and published multiword concepts across six perceptual modalities "
        "and five action effectors.",
        "Higher source ratings indicate stronger context-free normative "
        "associations with a sensory modality or bodily action. They identify "
        "lexical affordances for interpretation, not imagery guaranteed by the "
        "poem or an individual reader's experience.",
        (
            "Unmatched items remain missing rather than receiving a neutral or zero rating.",
            "Coverage, weighting, stopword treatment, and eligible-token counts should accompany comparisons.",
        ),
    ),
    "pronunciation": NarrativeReportProfile(
        "Pronunciation and Stress Report",
        "This report summarizes dictionary-supported syllable and stress evidence "
        "for the analyzed text.",
        "Pronunciation results reflect the selected dictionary entry and any local "
        "overrides; poetic performance may legitimately differ.",
        ("Unresolved or ambiguous pronunciations reduce coverage.",),
    ),
    "meter": NarrativeReportProfile(
        "Meter and Scansion Report",
        "This report summarizes candidate metrical patterns and their alignment "
        "with the pronunciation-supported stress evidence.",
        "Meter labels are ranked analytical candidates, not declarations of a "
        "single correct performance.",
        (
            "Consult line-level candidates and alignment operations in the companion CSV files.",
            "Performance-aware alternatives are reported only when enabled and supported.",
        ),
    ),
    "phonology": NarrativeReportProfile(
        "Rhyme and Sound Pattern Report",
        "This report summarizes end rhyme, internal rhyme, alliteration, "
        "assonance, consonance, and pronunciation coverage.",
        "Sound-pattern classifications depend on dictionary pronunciations and "
        "the configured evidence thresholds.",
        ("Unresolved pronunciations can hide sound relationships.",),
    ),
    "lexical_style": NarrativeReportProfile(
        "Lexical Style Report",
        "This report summarizes lexical diversity, word length, line word count, "
        "and stanza word count.",
        "Diversity measures respond differently to text length and token order; "
        "compare like with like and retain the configured parameters.",
        ("A missing value means the configured calculation was unavailable.",),
    ),
    "poetry_id": NarrativeReportProfile(
        "PoetryID Report",
        "This report summarizes descriptive evidence from the PoetryID feature set.",
        "The reported features are aids to inspection and comparison, not a "
        "judgment of poetic identity, authorship, quality, or genre.",
        ("Interpret every value with its denominator and method information.",),
    ),
    "inherited_form": NarrativeReportProfile(
        "Inherited Form Analysis Report",
        "This report compares the poem with ten versioned, source-backed "
        "inherited-form profiles using line, stanza, rhyme, meter, syllable, "
        "refrain, and end-word evidence when available.",
        "The leading result is a potential structural match. Consistency and "
        "confidence are transparent rule-based indices, not probabilities or "
        "declarations of genre identity, quality, historical intention, or tradition.",
        (
            "Missing pronunciation, meter, or rhyme evidence remains missing and lowers evidence coverage.",
            "Traditional definitions describe conventions that admit historical, linguistic, and artistic variation.",
            "Volta, kigo, semantic autonomy, and other interpretive features are not guessed when they cannot be scored defensibly.",
        ),
    ),
    "versemap": NarrativeReportProfile(
        "VerseMap Comparative Profile Report",
        "This report positions the analyzed poem relative to the versioned "
        "VerseMap public-domain reference corpus under Standard Profile 1.0.",
        "Nearby poems and poet centroids have smaller weighted standardized "
        "distances across shared evidence. The two plotted axes are PCA "
        "composites, not independently named literary traits.",
        (
            "Similarity is descriptive and is not evidence of authorship, influence, quality, or meaning.",
            "Missing evidence remains missing and reduces shared-evidence coverage rather than receiving a neutral score.",
            "Nearest-neighbor ranking uses the full registered feature space, not only the two-dimensional display.",
        ),
    ),
    "lexicon_explorer": NarrativeReportProfile(
        "Lexicon Explorer Report",
        "This report records one local word-or-phrase lookup across the installed "
        "VerseVAD lexical and pronunciation resources.",
        "The values are decontextualized source ratings, associations, corpus "
        "frequency evidence, retrospective norms, or dictionary pronunciation "
        "candidates. They support inspection rather than determine contextual "
        "meaning.",
        (
            "Missing or unavailable evidence remains missing rather than receiving a neutral value.",
            "Normalized VAD comparisons do not make source samples interchangeable.",
            "Pronunciation entries are dictionary candidates, not definitive poetic performances.",
        ),
    ),
    "phase2": NarrativeReportProfile(
        "VerseVAD Analysis Report",
        "This report brings together the principal results available for the "
        "analyzed text. Detailed audit records remain in the companion CSV files.",
        "The report describes computational evidence. It should support, rather "
        "than replace, close reading and documented scholarly judgment.",
        (
            "Coverage varies by lexicon and pronunciation resource.",
            "Missing lexical matches remain missing rather than being treated as neutral.",
        ),
    ),
    "corpus": NarrativeReportProfile(
        "VerseVAD Corpus Report",
        "This report summarizes the works and compatible result records represented "
        "in the selected VerseVAD project.",
        "Corpus comparisons are descriptive and are most defensible when the works "
        "share compatible configurations, resources, and coverage.",
        ("Use the companion CSV files for complete work-level and audit data.",),
    ),
    "compare_poems": NarrativeReportProfile(
        "VerseVAD Contrastive Evaluation Report",
        "This report places two poems side by side under one shared analytical "
        "configuration and records both values, B minus A differences, "
        "denominators, and coverage.",
        "The differences are descriptive prompts for contrastive close reading. "
        "They do not rank the poems, establish statistical significance, or "
        "reduce either poem to a single score.",
        (
            "Interpret only like-for-like source scales and methods.",
            "Coverage and eligible-token counts can differ between poems even when the configuration is shared.",
            "Missing evidence remains missing rather than receiving a neutral value.",
        ),
    ),
}


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_table_width(table, widths: Sequence[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, cell_width in zip(row.cells, widths, strict=True):
            cell.width = Inches(cell_width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=80, start=120, bottom=80, end=120)


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Title", 20, _DARK_BLUE, 0, 8),
        ("Heading 1", 16, _BLUE, 16, 8),
        ("Heading 2", 13, _BLUE, 12, 6),
        ("Heading 3", 12, _DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    _add_page_field(section.footer.paragraphs[0])


def _read_summary_rows(summary_csv: bytes) -> list[dict[str, str]]:
    text = summary_csv.decode("utf-8-sig")
    return [dict(row) for row in csv.DictReader(io.StringIO(text))]


def _display_name(value: str) -> str:
    return value.replace("_", " ").strip().title()


def _clean(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return "" if text.lower() in {"none", "nan"} else text


def _metric_sentence(row: Mapping[str, str]) -> str:
    metric = _display_name(_clean(row.get("metric", "Metric")))
    value = _clean(row.get("value", "")) or "not available"
    unit = _clean(row.get("unit_or_scale", row.get("unit", "")))
    denominator = _clean(row.get("denominator", ""))
    note = _clean(row.get("note", row.get("plain_language_note", "")))
    sentence = f"{metric}: {value}"
    if unit:
        sentence += f" ({unit})"
    sentence += "."
    if denominator:
        sentence += f" Denominator: {denominator}."
    if note:
        sentence += f" {note}"
    return sentence


def _add_metadata_table(
    document: Document,
    metadata: Sequence[tuple[str, str]],
) -> None:
    rows = [(label, value) for label, value in metadata if _clean(value)]
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
        _set_cell_shading(table.cell(index, 0), _PALE_GRAY)
        table.cell(index, 0).paragraphs[0].runs[0].bold = True
    _set_table_width(table, (2300, 7060))


def _add_companion_table(document: Document, filenames: Sequence[str]) -> None:
    if not filenames:
        return
    document.add_heading("Companion data files", level=1)
    document.add_paragraph(
        "The Word report is a readable orientation. These CSV files preserve the "
        "tabular values, denominators, provenance, and audit evidence:"
    )
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "CSV file"
    table.cell(0, 1).text = "Purpose"
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, _BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    _set_repeat_table_header(table.rows[0])
    for filename in filenames:
        cells = table.add_row().cells
        cells[0].text = filename
        stem = filename.removesuffix(".csv").replace("_", " ")
        cells[1].text = f"Machine-readable {stem} data."
    _set_table_width(table, (3900, 5460))


def _normalize_docx(package: bytes) -> bytes:
    """Return a byte-for-byte stable DOCX ZIP package."""

    source = io.BytesIO(package)
    target = io.BytesIO()
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(
        target,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as output:
        for name in sorted(archive.namelist()):
            information = zipfile.ZipInfo(name, _FIXED_TIMESTAMP)
            information.compress_type = zipfile.ZIP_DEFLATED
            information.external_attr = archive.getinfo(name).external_attr
            output.writestr(information, archive.read(name))
    return target.getvalue()


def _set_row_cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _set_cell_border(cell, *, color: str = "D7E0E6", size: str = "4") -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "start", "bottom", "end"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def _configure_comprehensive_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.70)
    section.right_margin = Inches(0.70)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04

    for name, size, color, before, after in (
        ("Title", 23, _DARK_BLUE, 0, 7),
        ("Heading 1", 15, _DARK_BLUE, 12, 6),
        ("Heading 2", 12, _BLUE, 9, 4),
        ("Heading 3", 10, _SLATE_BLUE, 7, 3),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "VERSEVAD  /  COMPUTATIONAL POETICS ANALYSIS"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(_SLATE_BLUE)
    _add_page_field(section.footer.paragraphs[0])


def _read_csv_table(content: bytes) -> tuple[list[str], list[dict[str, str]]]:
    text = content.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    return list(reader.fieldnames or ()), [dict(row) for row in reader]


def _numeric(value: object) -> float | None:
    text = _clean(value).replace(",", "")
    if not text or text.endswith("%"):
        return None
    try:
        number = float(text)
    except ValueError:
        return None
    return number if isfinite(number) else None


def _format_report_value(value: object, field: str = "") -> str:
    text = _clean(value)
    if not text:
        return "Not available"
    number = _numeric(text)
    if number is None:
        return text
    lowered = field.lower()
    count_fields = {
        "observation_count",
        "eligible_token_count",
        "matched_token_count",
        "unmatched_token_count",
        "eligible_type_count",
        "matched_type_count",
        "unmatched_type_count",
        "excluded_stopword_count",
        "excluded_non_content_count",
        "phrase_match_count",
        "matched_observations",
        "vad_matched_observations",
        "concreteness_matched_tokens",
        "line_number",
        "stanza_number",
        "rank",
        "version_number",
    }
    count_suffixes = (
        "_count",
        "_tokens",
        "_types",
        "_poems",
        "_lines",
        "_stanzas",
        "_sentences",
    )
    if number.is_integer() and (
        lowered in count_fields or lowered.endswith(count_suffixes)
    ):
        return f"{int(number):,}"
    if any(marker in lowered for marker in ("coverage", "proportion", "percent")):
        if 0 <= number <= 1:
            return f"{number * 100:.1f}%"
    return f"{number:.3f}"


def _friendly_header(value: str) -> str:
    replacements = {
        "id": "ID",
        "vad": "VAD",
        "aoa": "AoA",
        "sd": "SD",
        "iqr": "IQR",
        "pca": "PCA",
        "zipf": "Zipf",
        "sha256": "SHA-256",
    }
    words = value.replace("-", "_").split("_")
    return " ".join(replacements.get(word.lower(), word.title()) for word in words)


def _add_report_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    *,
    widths: Sequence[int] | None = None,
    font_size: float = 8.0,
) -> None:
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = str(header)
        _set_cell_shading(cell, _BLUE)
        _set_cell_border(cell, color=_BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.name = "Aptos"
            run.font.size = Pt(font_size)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    _set_repeat_table_header(table.rows[0])
    _set_row_cant_split(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = str(value)
            if row_index % 2:
                _set_cell_shading(cell, _LIGHTER_BLUE)
            _set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(font_size)
        _set_row_cant_split(table.rows[-1])
    if widths is None:
        usable = 10200
        widths = tuple(usable // len(headers) for _ in headers)
    _set_table_width(table, widths)


def _add_callout(document: Document, heading: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, _PALE_BLUE)
    _set_cell_border(cell, color="B7CDD9", size="6")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{heading}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(_DARK_BLUE)
    paragraph.add_run(body)
    _set_table_width(table, (10200,))


def _profile_family(row: Mapping[str, str]) -> str:
    return _PROFILE_FAMILY_BY_MODULE.get(_clean(row.get("module_id", "")), "")


def _file_family(filename: str) -> str:
    lowered = filename.lower()
    for hints, family in _FILE_FAMILY_HINTS:
        if any(hint in lowered for hint in hints):
            return family
    return ""


def _is_atomic_file(filename: str, row_count: int) -> bool:
    lowered = filename.lower()
    return row_count > 250 or any(marker in lowered for marker in _ATOMIC_FILE_MARKERS)


def _dataset_label(filename: str) -> str:
    return _friendly_header(filename.removesuffix(".csv"))


def _add_csv_dataset(
    document: Document,
    filename: str,
    fields: Sequence[str],
    records: Sequence[Mapping[str, str]],
) -> None:
    selected_fields = _REPORT_COLUMN_WHITELISTS.get(filename)
    if selected_fields:
        fields = [field for field in selected_fields if field in fields]
    if not fields:
        return
    document.add_heading(_dataset_label(filename), level=3)
    if not records:
        document.add_paragraph("No rows were available for this dataset.")
        return
    document.add_paragraph(
        f"{len(records):,} row(s). Numeric values are rounded to three decimal places in this report; companion CSV data retain full precision.",
    )
    identifiers = [
        field
        for field in fields
        if field.lower()
        in {
            "source",
            "metric",
            "section",
            "profile",
            "profile_id",
            "dimension",
            "category",
            "label",
            "candidate",
            "name",
            "poem",
            "poem_title",
            "title",
            "work",
        }
    ][:3]
    value_fields = [field for field in fields if field not in identifiers]
    if not value_fields:
        value_fields = list(fields)
        identifiers = []
    chunks = [value_fields[index : index + 5] for index in range(0, len(value_fields), 5)]
    for chunk_index, chunk in enumerate(chunks, start=1):
        columns = [*identifiers, *chunk]
        if len(chunks) > 1:
            caption = document.add_paragraph(
                f"Column group {chunk_index} of {len(chunks)}",
            )
            caption.runs[0].italic = True
            caption.runs[0].font.color.rgb = RGBColor.from_string(_SLATE_BLUE)
        table_rows = [
            tuple(_format_report_value(record.get(field, ""), field) for field in columns)
            for record in records
        ]
        _add_report_table(
            document,
            tuple(_friendly_header(field) for field in columns),
            table_rows,
            font_size=7.4 if len(columns) > 6 else 8.0,
        )


def _add_all_profile_matrix(
    document: Document,
    records: Sequence[Mapping[str, str]],
    *,
    companion_filename: str = "profile_metrics_all_compatible.csv",
) -> None:
    profiles = list(
        dict.fromkeys(_clean(row.get("profile_id", "")) for row in records)
    )
    profiles = [profile for profile in profiles if profile]
    profile_labels = {
        _clean(row.get("profile_id", "")): _clean(
            row.get("profile_label", row.get("profile_id", ""))
        )
        for row in records
    }
    metric_keys = list(
        dict.fromkeys(
            (
                _clean(row.get("source", row.get("source_id", ""))),
                _clean(row.get("metric", row.get("metric_id", ""))),
            )
            for row in records
        )
    )
    lookup = {
        (
            _clean(row.get("source", row.get("source_id", ""))),
            _clean(row.get("metric", row.get("metric_id", ""))),
            _clean(row.get("profile_id", "")),
        ): row
        for row in records
    }
    rows = []
    for source, metric in metric_keys:
        rows.append(
            (
                source,
                metric,
                *(
                    _format_report_value(
                        lookup.get((source, metric, profile), {}).get("value", ""),
                        "value",
                    )
                    for profile in profiles
                ),
            )
        )
    _add_report_table(
        document,
        ("Source", "Metric", *(profile_labels.get(profile, profile) for profile in profiles)),
        rows,
        font_size=7.1,
    )
    document.add_paragraph(
        "This matrix reports the primary value for every compatible lexical profile. "
        f"The companion {companion_filename} retains full-precision "
        "statistics, counts, exclusions, and coverage. Statistical fields remain blank "
        "when they do not apply; cumulative fields are populated only for metric "
        "families with a defined additive interpretation."
    )


def _profile_display_rows(
    records: Sequence[Mapping[str, str]],
) -> tuple[
    list[tuple[str, ...]],
    list[tuple[str, ...]],
    list[tuple[str, ...]],
]:
    from versevad.metric_capabilities import metric_capabilities

    primary: list[tuple[str, ...]] = []
    distribution: list[tuple[str, ...]] = []
    cumulative: list[tuple[str, ...]] = []
    for row in records:
        source = _clean(row.get("source", row.get("source_id", "")))
        metric = _clean(row.get("metric", row.get("metric_id", "")))
        profile = _clean(row.get("profile_label", row.get("profile_id", "")))
        primary.append(
            (
                source,
                metric,
                profile,
                _format_report_value(row.get("value", ""), "value"),
                _format_report_value(row.get("median", ""), "median"),
                _format_report_value(row.get("observation_count", ""), "observation_count"),
                _format_report_value(row.get("token_coverage", ""), "token_coverage"),
                _clean(row.get("unit", "")),
            )
        )
        q1 = _format_report_value(row.get("first_quartile", ""), "first_quartile")
        q3 = _format_report_value(row.get("third_quartile", ""), "third_quartile")
        minimum = _format_report_value(row.get("minimum", ""), "minimum")
        maximum = _format_report_value(row.get("maximum", ""), "maximum")
        def display_range(first: str, second: str) -> str:
            if first == "Not available" and second == "Not available":
                return "Not available"
            if first == "Not available":
                return second
            if second == "Not available":
                return first
            return f"{first} to {second}"

        capability = metric_capabilities(_clean(row.get("module_id", "")))
        try:
            observation_count = int(float(_clean(row.get("observation_count", "0")) or 0))
        except ValueError:
            observation_count = 0
        if capability.supports_dispersion and observation_count >= 2:
            distribution.append(
                (
                    source,
                    metric,
                    profile,
                    _format_report_value(
                        row.get("population_standard_deviation", ""),
                        "population_standard_deviation",
                    ),
                    display_range(q1, q3),
                    display_range(minimum, maximum),
                    _format_report_value(
                        row.get("average_deviation_from_mean", ""),
                        "average_deviation_from_mean",
                    ),
                )
            )
        if capability.supports_raw_accumulation and _clean(row.get("cumulative_value", "")):
            cumulative.append(
                (
                    source,
                    metric,
                    profile,
                    _format_report_value(row.get("cumulative_value", ""), "cumulative_value"),
                    _clean(row.get("unit", "")),
                    _format_report_value(row.get("observation_count", ""), "observation_count"),
                )
            )
    return primary, distribution, cumulative


def _add_profile_family_tables(
    document: Document,
    records: Sequence[Mapping[str, str]],
) -> None:
    primary, distribution, cumulative = _profile_display_rows(records)
    document.add_heading("Primary values and coverage", level=2)
    _add_report_table(
        document,
        ("Source", "Metric", "Profile", "Primary value", "Median", "N", "Coverage", "Unit"),
        primary,
        font_size=7.3,
    )
    if distribution:
        document.add_heading("Within-text dispersion and volatility", level=2)
        _add_report_table(
            document,
            ("Source", "Metric", "Profile", "Population SD", "IQR", "Min to Max", "Mean Abs. Deviation"),
            distribution,
            font_size=7.2,
        )
        document.add_paragraph(
            "Dispersion is reported only for continuous evidence with at least two "
            "contributing observations. Categorical associations and one-observation "
            "summaries are omitted rather than presented as zero variability."
        )
    if cumulative:
        has_intensity = any(
            _clean(row.get("module_id", "")) == "emotion_intensity"
            and _clean(row.get("cumulative_value", ""))
            for row in records
        )
        document.add_heading(
            "Cumulative Emotion Intensity Load"
            if has_intensity
            else "Method-defined cumulative lexical magnitude",
            level=2,
        )
        _add_report_table(
            document,
            ("Source", "Metric", "Profile", "Cumulative value", "Source unit", "Observations"),
            cumulative,
            font_size=7.2,
        )
        document.add_paragraph(
            "Cumulative values appear only where the metric methodology defines an "
            "additive magnitude. VerseVAD does not construct generic AoA, Zipf, "
            "concreteness, or word-length totals."
        )
    if any(_clean(row.get("absolute_midpoint_load", "")) for row in records):
        document.add_heading("Midpoint-relative lexical load", level=2)
        midpoint_rows = [
            (
                _clean(row.get("metric", row.get("metric_id", ""))),
                _clean(row.get("profile_label", row.get("profile_id", ""))),
                _format_report_value(row.get("above_midpoint_load", ""), "above_midpoint_load"),
                _format_report_value(row.get("below_midpoint_load", ""), "below_midpoint_load"),
                _format_report_value(row.get("net_midpoint_load", ""), "net_midpoint_load"),
                _format_report_value(row.get("absolute_midpoint_load", ""), "absolute_midpoint_load"),
                _format_report_value(
                    (
                        float(row.get("above_midpoint_load", ""))
                        / float(row.get("observation_count", "")) * 100
                    )
                    if _clean(row.get("above_midpoint_load", ""))
                    and float(row.get("observation_count", "0") or 0) > 0
                    else "",
                    "above_midpoint_load_per_100",
                ),
                _format_report_value(
                    (
                        float(row.get("below_midpoint_load", ""))
                        / float(row.get("observation_count", "")) * 100
                    )
                    if _clean(row.get("below_midpoint_load", ""))
                    and float(row.get("observation_count", "0") or 0) > 0
                    else "",
                    "below_midpoint_load_per_100",
                ),
                _format_report_value(
                    (
                        float(row.get("net_midpoint_load", ""))
                        / float(row.get("observation_count", "")) * 100
                    )
                    if _clean(row.get("net_midpoint_load", ""))
                    and float(row.get("observation_count", "0") or 0) > 0
                    else "",
                    "net_midpoint_load_per_100",
                ),
                _format_report_value(
                    (
                        float(row.get("absolute_midpoint_load", ""))
                        / float(row.get("observation_count", "")) * 100
                    )
                    if _clean(row.get("absolute_midpoint_load", ""))
                    and float(row.get("observation_count", "0") or 0) > 0
                    else "",
                    "absolute_midpoint_load_per_100",
                ),
                (
                    "per 100 matched tokens"
                    if _clean(row.get("weighting", "")) == "TOKEN"
                    else "per 100 matched types"
                ),
            )
            for row in records
            if _clean(row.get("absolute_midpoint_load", ""))
        ]
        _add_report_table(
            document,
            (
                "Metric", "Profile", "Above", "Below", "Net", "Absolute",
                "Above / 100", "Below / 100", "Net / 100", "Absolute / 100",
                "Normalized denominator",
            ),
            midpoint_rows,
            font_size=6.8,
        )
        document.add_paragraph(
            "Raw rating totals are not reader-facing results. Midpoint-relative loads "
            "sum distance from 0.5; normalized values use the matched-token denominator "
            "for token weighting and matched-type denominator for type weighting."
        )


def _profile_dashboard_rows(
    records: Sequence[Mapping[str, str]],
) -> list[tuple[str, str, str, str]]:
    preferred = [
        row
        for row in records
        if _clean(row.get("profile_id", ""))
        in {
            "STOPWORD_EXCLUDED__TOKEN",
            "stopword_excluded-token_weighted",
            "stopword_excluded_token_weighted",
        }
    ]
    source = preferred or list(records)
    rows: list[tuple[str, str, str, str]] = []
    seen: set[tuple[str, str]] = set()
    for row in source:
        key = (_clean(row.get("source_id", "")), _clean(row.get("metric_id", "")))
        if key in seen:
            continue
        seen.add(key)
        rows.append(
            (
                _clean(row.get("source", row.get("source_id", ""))),
                _clean(row.get("metric", row.get("metric_id", ""))),
                _format_report_value(row.get("value", ""), "value"),
                _clean(row.get("unit", "")),
            )
        )
        if len(rows) >= 14:
            break
    return rows


def build_comprehensive_analysis_report(
    *,
    export_files: Mapping[str, bytes],
    text_title: str = "",
    author: str = "",
    analysis_timestamp: str = "",
    export_mode: str = "complete_audit",
    visible_section: str = "",
    workspace_label: str = "Single Poem",
    text_id: str = "",
    result_id: str = "",
    source_sha256: str = "",
    analysis_profiles: Sequence[str] = (),
    active_preset: str = "",
    source_notes: str = "",
    software_version: str = "",
    warnings: Sequence[str] = (),
    resources: Sequence[str] = (),
    methods_reproducibility: Sequence[str] = (),
    calculated_modules: Sequence[str] = (),
) -> bytes:
    """Build the full readable analysis report directly from exported evidence."""

    if export_mode not in {"current_view", "complete_audit"}:
        raise ValueError(f"Unknown export mode: {export_mode}")
    parsed: OrderedDict[str, tuple[list[str], list[dict[str, str]]]] = OrderedDict()
    for filename in sorted(export_files):
        if filename.lower().endswith(".csv"):
            try:
                parsed[filename] = _read_csv_table(export_files[filename])
            except (UnicodeDecodeError, csv.Error):
                continue

    profile_records = parsed.get("profile_metrics_selected.csv", ([], []))[1]
    family_profiles: dict[str, list[dict[str, str]]] = {
        family.family_id: [] for family in COMPREHENSIVE_REPORT_FAMILIES
    }
    for row in profile_records:
        family = _profile_family(row)
        if family:
            family_profiles[family].append(row)

    family_files: dict[str, list[str]] = {
        family.family_id: [] for family in COMPREHENSIVE_REPORT_FAMILIES
    }
    calculated_families = {
        family
        for module in calculated_modules
        if (family := _CALCULATED_FAMILY_BY_MODULE.get(_clean(module)))
    }
    excluded_inventory: list[tuple[str, str, str]] = []
    special_files = {
        "profile_metrics_selected.csv",
        "profile_metrics_all_compatible.csv",
        "scholar_summary.csv",
        "csv_reading_guide.csv",
        "phase2_manifest.csv",
    }
    for filename, (_fields, records) in parsed.items():
        if filename in special_files:
            continue
        family = _file_family(filename)
        if not family:
            continue
        if family_profiles[family] and filename not in _PROFILE_BACKED_REPORT_FILES:
            excluded_inventory.append(
                (
                    filename,
                    f"{len(records):,}",
                    "Supporting aggregate detail retained in the companion bundle; selected-profile results are reported in the family tables",
                )
            )
            continue
        if _is_atomic_file(filename, len(records)):
            excluded_inventory.append(
                (filename, f"{len(records):,}", "Atomic or high-volume evidence retained in the companion bundle")
            )
        else:
            family_files[family].append(filename)

    document = Document()
    _configure_comprehensive_document(document)
    report_kind = "Current View Report" if export_mode == "current_view" else "Complete Audit Report"
    document.core_properties.title = "VerseVAD Computational Poetics Analysis Report"
    document.core_properties.subject = report_kind
    document.core_properties.author = "VerseVAD"
    document.core_properties.created = _FIXED_CORE_DATE
    document.core_properties.modified = _FIXED_CORE_DATE

    kicker = document.add_paragraph("VERSEVAD  /  DIGITAL HUMANITIES FOR EVERYONE")
    kicker.runs[0].bold = True
    kicker.runs[0].font.size = Pt(8.5)
    kicker.runs[0].font.color.rgb = RGBColor.from_string(_SLATE_BLUE)
    title = document.add_paragraph(style="Title")
    title.add_run("Computational Poetics\nAnalysis Report")
    subtitle = document.add_paragraph(report_kind)
    subtitle.style = document.styles["Subtitle"]
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(_BLUE)
    _add_callout(
        document,
        "Interpretive principle",
        "This report organizes computational evidence for close reading. It does not reduce a poem to a score, determine meaning, or replace documented scholarly judgment.",
    )
    _add_metadata_table(
        document,
        (
            ("Text", text_title or "Untitled text"),
            ("Author / creator", author or "[Enter author or creator]"),
            ("Analyst", "[Enter analyst name]"),
            ("Research question", "[Enter research question]"),
            ("Workspace", workspace_label),
            ("Report scope", visible_section or report_kind),
            ("Analysis date and time", analysis_timestamp or "Not recorded"),
            ("VerseVAD version", software_version or "Not recorded"),
            ("Analysis profile", ", ".join(analysis_profiles) or "Not recorded"),
            ("Module preset", active_preset or "Custom"),
            ("Source / bibliographic notes", source_notes),
        ),
    )
    document.add_paragraph()
    document.add_paragraph("Prepared for transparent, reproducible computational close reading.")
    document.add_page_break()

    document.add_heading("1. How to Read This Report", level=1)
    document.add_paragraph(
        "Means are the primary location measures for continuous lexical metrics. Medians, dispersion, cumulative load, coverage, and denominators are shown where the underlying analysis supplies them. Values are rounded to three decimal places here; companion CSV files retain the original precision."
    )
    if export_mode == "current_view":
        document.add_paragraph(
            "This Current View report includes the selected report family and the globally selected lexical scope/weighting profiles. Calculated families outside the selected report section are identified separately from modules that were not calculated."
        )
    else:
        document.add_paragraph(
            "This Complete Audit report includes every calculated aggregate available in the completed analysis. Disabled or unavailable modules are marked Not calculated. Atomic token-, line-, pair-, and operation-level evidence remains in the companion audit bundle so this document stays readable."
        )
    _add_callout(
        document,
        "Rounding",
        "Report values use three decimal places where appropriate. Use the CSV audit files for full-precision computation or secondary analysis.",
    )

    document.add_heading("2. Executive Metric Dashboard", level=1)
    dashboard = _profile_dashboard_rows(profile_records)
    if dashboard:
        _add_report_table(document, ("Source", "Metric", "Mean", "Unit"), dashboard)
    else:
        document.add_paragraph("No compatible lexical-profile rows were available for the dashboard.")

    document.add_heading("3. Coverage and Evidence Quality", level=1)
    coverage_rows: list[tuple[str, ...]] = []
    seen_coverage: set[tuple[str, str, str]] = set()
    for row in profile_records:
        source = _clean(row.get("source", row.get("source_id", "")))
        module = _friendly_header(_clean(row.get("module_id", "")))
        profile = _clean(row.get("profile_label", row.get("profile_id", "")))
        key = (source, module, profile)
        if key in seen_coverage:
            continue
        seen_coverage.add(key)
        weighting = _clean(row.get("weighting", ""))
        if weighting.casefold() in {"type", "type_weighted", "type-weighted"}:
            eligible = row.get("eligible_type_count", "")
            matched = row.get("matched_type_count", "")
            coverage = row.get("type_coverage", "")
            denominator = "eligible types"
        else:
            eligible = row.get("eligible_token_count", "")
            matched = row.get("matched_token_count", "")
            coverage = row.get("token_coverage", "")
            denominator = "eligible tokens"
        coverage_rows.append(
            (
                source,
                module,
                profile,
                _format_report_value(eligible, "eligible_token_count"),
                _format_report_value(matched, "matched_token_count"),
                _format_report_value(coverage, "coverage"),
                denominator,
            )
        )
    if coverage_rows:
        _add_report_table(
            document,
            ("Resource", "Module", "Profile", "Eligible", "Matched", "Coverage", "Denominator"),
            coverage_rows,
            font_size=7.2,
        )
    else:
        document.add_paragraph("No profile-aware coverage records were available.")
    document.add_paragraph(
        "Coverage is part of the result. Token weighting uses occurrence counts; "
        "type weighting uses distinct matched lexical identities. Unmatched evidence "
        "remains missing rather than receiving a neutral value."
    )

    document.add_heading("4. Module Reporting Status", level=1)
    status_rows = []
    first_reported_family = True
    for family in COMPREHENSIVE_REPORT_FAMILIES:
        present = bool(family_profiles[family.family_id] or family_files[family.family_id])
        if present:
            status = "Reported"
            note = "Calculated evidence is included below."
        elif family.family_id in calculated_families and export_mode == "current_view":
            status = "Calculated, not included"
            selected_section = visible_section or "Current View"
            note = f"Calculated, but not included in the selected {selected_section} report section."
        elif family.family_id in calculated_families:
            status = "Calculated; companion data only"
            note = "Calculated evidence is retained in the companion audit files."
        else:
            status = "Not calculated"
            note = "Module disabled, unavailable, or unsupported for this text."
        status_rows.append((family.title, status, note))
    _add_report_table(document, ("Metric family", "Status", "Reason"), status_rows)

    section_number = 5
    for family in COMPREHENSIVE_REPORT_FAMILIES:
        profile_rows = family_profiles[family.family_id]
        dataset_files = family_files[family.family_id]
        if not profile_rows and not dataset_files:
            continue
        if not first_reported_family:
            document.add_page_break()
        first_reported_family = False
        document.add_heading(f"{section_number}. {family.title}", level=1)
        section_number += 1
        document.add_paragraph(family.explanation)
        for caution in family.cautions:
            _add_callout(document, "Interpretive caution", caution)
        if profile_rows:
            _add_profile_family_tables(document, profile_rows)
            if family.family_id == "vad" and len(
                {_clean(row.get("source_id", "")) for row in profile_rows}
            ) > 1:
                document.add_heading("Cross-Lexicon VAD Comparison", level=2)
                document.add_paragraph(
                    "The resource-specific values above remain primary. This comparison "
                    "places active VAD sources beside one another without averaging them. "
                    "Differences can reflect norming samples, coverage, scales, and source "
                    "vocabularies rather than disagreement about a single true score."
                )
                comparison_rows = [
                    (
                        _clean(row.get("source", row.get("source_id", ""))),
                        _clean(row.get("metric", row.get("metric_id", ""))),
                        _clean(row.get("profile_label", row.get("profile_id", ""))),
                        _format_report_value(row.get("value", ""), "value"),
                        _format_report_value(row.get("token_coverage", ""), "token_coverage"),
                    )
                    for row in profile_rows
                ]
                _add_report_table(
                    document,
                    ("Resource", "Dimension", "Profile", "Normalized value", "Token coverage"),
                    comparison_rows,
                    font_size=7.2,
                )
        for filename in dataset_files:
            fields, records = parsed[filename]
            _add_csv_dataset(document, filename, fields, records)

    all_profile_fields, all_profile_rows = parsed.get(
        "profile_metrics_all_compatible.csv", ([], [])
    )
    if export_mode == "complete_audit" and all_profile_rows:
        document.add_page_break()
        corpus_profile_matrix = "corpus_vad_profiles.csv" in parsed
        matrix_title = (
            "All Compatible VAD Profiles"
            if corpus_profile_matrix
            else "All Compatible Lexical Profiles"
        )
        document.add_heading(f"{section_number}. {matrix_title}", level=1)
        section_number += 1
        document.add_paragraph(
            "This appendix records every calculated lexical scope and weighting combination. The main report emphasizes the selected profiles; this table preserves the complete compatible profile space."
        )
        _add_all_profile_matrix(
            document,
            all_profile_rows,
            companion_filename=(
                "corpus_vad_profiles.csv"
                if corpus_profile_matrix
                else "profile_metrics_all_compatible.csv"
            ),
        )

    document.add_page_break()
    document.add_heading(f"{section_number}. Warnings and Limitations", level=1)
    section_number += 1
    unique_warnings = [warning for warning in dict.fromkeys(_clean(item) for item in warnings) if warning]
    if unique_warnings:
        for warning in unique_warnings:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(warning)
    else:
        document.add_paragraph("No additional module warnings were recorded in this export.")
    document.add_paragraph(
        "Coverage is part of the result. Unmatched or source-unrated evidence remains missing rather than receiving a neutral value. Comparisons should retain lexical scope, weighting, eligible-observation counts, resource versions, and coverage."
    )

    document.add_heading(f"{section_number}. Methods and Reproducibility", level=1)
    section_number += 1
    _add_metadata_table(
        document,
        (
            ("Text ID", text_id),
            ("Result / text-version ID", result_id),
            ("Source-text SHA-256", source_sha256),
            ("Report mode", export_mode),
            ("Selected report section", visible_section),
        ),
    )
    for paragraph_text in methods_reproducibility:
        if _clean(paragraph_text):
            document.add_paragraph(_clean(paragraph_text))
    if resources:
        document.add_heading("Validated resources", level=2)
        for resource in resources:
            document.add_paragraph(_clean(resource), style="List Bullet")

    document.add_heading(f"{section_number}. Interpretive Synthesis", level=1)
    section_number += 1
    for label in (
        "Primary observations",
        "Patterns that support the research question",
        "Counterevidence, ambiguity, and surprises",
        "Close-reading passages to revisit",
    ):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        paragraph.add_run("[Enter analyst synthesis]")

    document.add_heading(f"{section_number}. Companion Audit Files", level=1)
    csv_names = [name for name in export_files if name.lower().endswith(".csv")]
    _add_companion_table(document, csv_names)
    if excluded_inventory:
        document.add_heading("High-volume evidence retained outside the report", level=2)
        _add_report_table(
            document,
            ("File", "Rows", "Purpose"),
            excluded_inventory,
        )
    document.add_paragraph(
        "Retain this Word report with its CSV files, manifests, file inventory, and reproducibility README. The DOCX is designed for reading; the audit bundle remains the authoritative machine-readable record."
    )

    output = io.BytesIO()
    document.save(output)
    return _normalize_docx(output.getvalue())


def build_narrative_report(
    *,
    profile: NarrativeReportProfile,
    summary_rows: Iterable[Mapping[str, str]],
    companion_csv_files: Sequence[str],
    text_title: str = "",
    text_id: str = "",
    result_id: str = "",
    warnings: Sequence[str] = (),
    additional_paragraphs: Sequence[str] = (),
    methods_reproducibility: Sequence[str] = (),
) -> bytes:
    """Build one accessible narrative report with stable package bytes."""

    document = Document()
    _configure_document(document)
    document.core_properties.title = profile.title
    document.core_properties.subject = "VerseVAD narrative analysis report"
    document.core_properties.author = "VerseVAD"
    document.core_properties.created = _FIXED_CORE_DATE
    document.core_properties.modified = _FIXED_CORE_DATE

    title = document.add_paragraph(style="Title")
    title.add_run(profile.title)
    strapline = document.add_paragraph("VerseVAD · Local analytical evidence")
    strapline.style = document.styles["Subtitle"]
    _add_metadata_table(
        document,
        (
            ("Text", text_title),
            ("Text ID", text_id),
            ("Result ID", result_id),
        ),
    )

    document.add_heading("Scope and interpretation", level=1)
    document.add_paragraph(profile.scope)
    document.add_paragraph(profile.interpretation)
    for paragraph in additional_paragraphs:
        if _clean(paragraph):
            document.add_paragraph(_clean(paragraph))

    rows = list(summary_rows)
    if rows:
        document.add_heading("Key findings", level=1)
        current_section = ""
        for row in rows:
            section = _clean(row.get("section", "")) or "summary"
            if section != current_section:
                document.add_heading(_display_name(section), level=2)
                current_section = section
            document.add_paragraph(_metric_sentence(row))
    else:
        document.add_heading("Key findings", level=1)
        document.add_paragraph("No summary rows were available for this report.")

    cautions = [*profile.cautions, *(_clean(item) for item in warnings)]
    cautions = [item for item in cautions if item]
    if cautions:
        document.add_heading("Coverage and cautions", level=1)
        for caution in cautions:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(caution)

    if methods_reproducibility:
        document.add_heading("Methods and Reproducibility", level=1)
        for paragraph_text in methods_reproducibility:
            if _clean(paragraph_text):
                document.add_paragraph(_clean(paragraph_text))

    _add_companion_table(document, companion_csv_files)
    document.add_heading("How to cite and reproduce", level=1)
    document.add_paragraph(
        "Retain this report with its companion CSV files. Record the VerseVAD "
        "version, configuration identifiers, resource names and hashes, and any "
        "local overrides shown in the exported data."
    )

    output = io.BytesIO()
    document.save(output)
    return _normalize_docx(output.getvalue())


def build_narrative_report_from_summary_csv(
    module_name: str,
    summary_csv: bytes,
    *,
    companion_csv_files: Sequence[str],
    text_title: str = "",
    text_id: str = "",
    result_id: str = "",
    warnings: Sequence[str] = (),
    additional_paragraphs: Sequence[str] = (),
    methods_reproducibility: Sequence[str] = (),
) -> bytes:
    """Build a report from any VerseVAD summary CSV using a named profile."""

    return build_narrative_report(
        profile=REPORT_PROFILES[module_name],
        summary_rows=_read_summary_rows(summary_csv),
        companion_csv_files=companion_csv_files,
        text_title=text_title,
        text_id=text_id,
        result_id=result_id,
        warnings=warnings,
        additional_paragraphs=additional_paragraphs,
        methods_reproducibility=methods_reproducibility,
    )
