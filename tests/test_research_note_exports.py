from __future__ import annotations

import io
import zipfile

from docx import Document

from versevad.exports.archive_contract import (
    COMPLETE_AUDIT_ANALYSIS_REPORT_PATH,
)
from versevad.exports.docx_report import (
    NarrativeReportProfile,
    build_narrative_report,
)
from versevad.exports.research_notes import (
    add_research_notes_to_audit_bundle,
    append_research_notes_to_docx,
    research_notes_csv,
)
from versevad.research_library import ResearchNote


def _note() -> ResearchNote:
    return ResearchNote(
        note_id="note-1",
        parent_type="analysis",
        parent_id="analysis-1",
        analysis_id="analysis-1",
        project_id="",
        module="Affective Evidence",
        metric="Valence",
        anchor_type="metric",
        anchor_label="Valence → token-weighted mean",
        title="Refrain question",
        body="Compare the low mean with the poem's repeated positive refrain.",
        tags=("repetition", "valence"),
        include_in_export=True,
        created_at="2026-07-29T12:00:00+00:00",
        updated_at="2026-07-29T12:05:00+00:00",
    )


def _report() -> bytes:
    return build_narrative_report(
        profile=NarrativeReportProfile(
            title="Test Report",
            scope="Scope.",
            interpretation="Interpretation.",
        ),
        summary_rows=(
            {
                "section": "Overview",
                "metric": "Matched",
                "value": "3",
            },
        ),
        companion_csv_files=("summary.csv",),
    )


def test_selected_notes_append_to_word_report_without_replacing_it() -> None:
    report = append_research_notes_to_docx(
        _report(),
        (_note(),),
        include_metadata=True,
    )
    paragraphs = [paragraph.text for paragraph in Document(io.BytesIO(report)).paragraphs]

    assert "Test Report" in paragraphs
    assert "Research Notes" in paragraphs
    assert "Refrain question" in paragraphs
    assert any("separate from VerseVAD" in paragraph for paragraph in paragraphs)
    assert any("note-1" in paragraph for paragraph in paragraphs)


def test_notes_csv_can_exclude_private_metadata() -> None:
    content = research_notes_csv((_note(),), include_metadata=False).decode(
        "utf-8-sig"
    )

    assert "title,body,anchor_type" in content
    assert "Refrain question" in content
    assert "note_id" not in content
    assert "analysis-1" not in content


def test_audit_bundle_adds_csv_markdown_and_note_inclusive_docx() -> None:
    source = io.BytesIO()
    with zipfile.ZipFile(source, "w") as archive:
        archive.writestr("VerseVAD_analysis_report.docx", _report())
        archive.writestr("summary.csv", b"metric,value\nMatched,3\n")

    updated = add_research_notes_to_audit_bundle(
        source.getvalue(),
        (_note(),),
        include_metadata=False,
    )

    with zipfile.ZipFile(io.BytesIO(updated)) as archive:
        assert sorted(archive.namelist()) == [
            "VerseVAD_analysis_report.docx",
            "research_notes.csv",
            "research_notes.md",
            "summary.csv",
        ]
        paragraphs = [
            paragraph.text
            for paragraph in Document(
                io.BytesIO(archive.read("VerseVAD_analysis_report.docx"))
            ).paragraphs
        ]
    assert "Research Notes" in paragraphs


def test_complete_audit_bundle_appends_notes_to_nested_report() -> None:
    source = io.BytesIO()
    with zipfile.ZipFile(source, "w") as archive:
        archive.writestr(COMPLETE_AUDIT_ANALYSIS_REPORT_PATH, _report())

    updated = add_research_notes_to_audit_bundle(
        source.getvalue(),
        (_note(),),
        include_metadata=False,
    )

    with zipfile.ZipFile(io.BytesIO(updated)) as archive:
        report = archive.read(COMPLETE_AUDIT_ANALYSIS_REPORT_PATH)
    paragraphs = [
        paragraph.text for paragraph in Document(io.BytesIO(report)).paragraphs
    ]
    assert "Research Notes" in paragraphs
