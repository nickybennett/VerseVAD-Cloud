from __future__ import annotations

import csv
import io
import zipfile

import pytest
from docx import Document

from versevad.db import CorpusMetricRecord, CorpusTextRecord, ProjectRecord
from versevad.exports.corpus_csv import build_corpus_export_bundle
from versevad.analysis_profiles import ProfileSelection


def test_corpus_export_contains_data_report_and_reproducibility_records() -> None:
    project = ProjectRecord(
        project_id="project-export",
        title="Export Project",
        description="Synthetic export fixture",
        researcher="Researcher",
        created_at="2026-07-25T00:00:00+00:00",
        updated_at="2026-07-25T00:00:00+00:00",
    )
    text = CorpusTextRecord(
        text_id="text-export",
        text_version_id="text-export:v1",
        project_id=project.project_id,
        title="Fixture",
        source_name="fixture.txt",
        relative_path="fixture.txt",
        author="",
        collection="",
        date_label="",
        genre="",
        notes="",
        custom_metadata={},
        original_text="The full text stays in the project database.",
        text_sha256="a" * 64,
        imported_at="2026-07-25T00:00:00+00:00",
        updated_at="2026-07-25T00:00:00+00:00",
    )

    first = build_corpus_export_bundle(project, (text,), (), ())
    second = build_corpus_export_bundle(project, (text,), (), ())

    assert first == second
    with zipfile.ZipFile(io.BytesIO(first)) as archive:
        names = set(archive.namelist())
        assert "01_REPORTS/Corpus_Report.docx" in names
        assert "01_REPORTS/Coverage_and_Data_Quality.docx" in names
        assert "03_MASTER_DATA/Master_Metrics.csv" in names
        assert "02_METRIC_TABLES/Corpus_Summary.csv" in names
        assert "03_MASTER_DATA/Selected_Profiles.csv" in names
        assert "03_MASTER_DATA/All_Profiles.csv" in names
        assert "03_MASTER_DATA/Scope_Token_Counts.csv" in names
        assert "03_MASTER_DATA/Works.csv" in names
        assert "05_REPRODUCIBILITY/Methodology.csv" in names
        assert "05_REPRODUCIBILITY/REPRODUCIBILITY_README.txt" in names
        assert "05_REPRODUCIBILITY/FILE_INVENTORY.csv" in names
        assert not any(name.endswith((".json", ".xlsx")) for name in names)
        rows = list(
            csv.DictReader(
                io.StringIO(
                    archive.read("03_MASTER_DATA/Works.csv").decode("utf-8-sig")
                )
            )
        )
        assert rows[0]["title"] == "Fixture"
        assert "original_text" not in rows[0]
        report = archive.read("01_REPORTS/Corpus_Report.docx")
        assert report.startswith(b"PK")
        document = Document(io.BytesIO(report))
        text_content = "\n".join(
            paragraph.text for paragraph in document.paragraphs
        )
        assert "Computational Poetics\nAnalysis Report" in text_content
        assert "How to Read This Report" in text_content
        assert "Companion Audit Files" in text_content


def test_corpus_export_reports_both_vad_dispersion_levels() -> None:
    project = ProjectRecord(
        project_id="project-dispersion",
        title="Dispersion Project",
        description="Synthetic dispersion fixture",
        researcher="Researcher",
        created_at="2026-07-26T00:00:00+00:00",
        updated_at="2026-07-26T00:00:00+00:00",
    )
    texts = tuple(
        CorpusTextRecord(
            text_id=text_id,
            text_version_id=f"{text_id}:v1",
            project_id=project.project_id,
            title=title,
            source_name=f"{text_id}.txt",
            relative_path=f"{text_id}.txt",
            author="",
            collection="",
            date_label="",
            genre="",
            notes="",
            custom_metadata={},
            original_text="Fixture.",
            text_sha256=character * 64,
            imported_at="2026-07-26T00:00:00+00:00",
            updated_at="2026-07-26T00:00:00+00:00",
        )
        for text_id, title, character in (
            ("first", "First", "a"),
            ("second", "Second", "b"),
        )
    )

    def metric(
        text: CorpusTextRecord,
        metric_name: str,
        value: float,
        observations: int,
    ) -> CorpusMetricRecord:
        return CorpusMetricRecord(
            run_id=f"run-{text.text_id}",
            text_id=text.text_id,
            text_version_id=text.text_version_id,
            title=text.title,
            author="",
            collection="",
            date_label="",
            genre="",
            lexicon_id="fixture-vad",
            lexicon="Fixture VAD",
            value_kind="vad",
            metric=metric_name,
            dimension="valence",
            category="",
            weighting="token",
            scale="normalized_0_1",
            denominator=f"{observations} included matched observations",
            value=value,
            observations=observations,
            matched_tokens=observations,
            lexical_tokens=observations,
            coverage=1.0,
            completed_at="2026-07-26T00:00:00+00:00",
            analysis_view="all_matched",
        )

    metrics = (
        metric(texts[0], "vad_mean", 0.2, 2),
        metric(texts[0], "vad_standard_deviation", 0.1, 2),
        metric(texts[1], "vad_mean", 0.8, 3),
        metric(texts[1], "vad_standard_deviation", 0.2, 3),
    )
    archive_bytes = build_corpus_export_bundle(
        project,
        texts,
        metrics,
        (),
    )

    with zipfile.ZipFile(io.BytesIO(archive_bytes)) as archive:
        profile_rows = list(
            csv.DictReader(
                io.StringIO(
                    archive.read("04_AUDIT/corpus_vad_profiles.csv").decode("utf-8-sig")
                )
            )
        )
        assert profile_rows[0][
            "pooled_lexical_rating_standard_deviation"
        ] != ""
        assert float(
            profile_rows[0]["poem_mean_standard_deviation"]
        ) == pytest.approx(0.3)
        report = Document(io.BytesIO(archive.read("01_REPORTS/Corpus_Report.docx")))
        report_text = "\n".join(
            [
                *(paragraph.text for paragraph in report.paragraphs),
                *(
                    cell.text
                    for table in report.tables
                    for row in table.rows
                    for cell in row.cells
                ),
            ]
        ).casefold()
        assert "valence pooled-observation mean" in report_text
        assert "valence equal-work mean" in report_text
        assert "population sd" in report_text


def test_corpus_export_repairs_legacy_type_metadata_and_reports_matched_count() -> None:
    project = ProjectRecord(
        project_id="project-type-metadata",
        title="Type Metadata Project",
        description="Synthetic legacy metadata fixture",
        researcher="Researcher",
        created_at="2026-08-10T00:00:00+00:00",
        updated_at="2026-08-10T00:00:00+00:00",
    )
    text = CorpusTextRecord(
        text_id="type-text",
        text_version_id="type-text:v1",
        project_id=project.project_id,
        title="Fixture",
        source_name="fixture.txt",
        relative_path="fixture.txt",
        author="",
        collection="",
        date_label="",
        genre="",
        notes="",
        custom_metadata={},
        original_text="Fixture.",
        text_sha256="c" * 64,
        imported_at="2026-08-10T00:00:00+00:00",
        updated_at="2026-08-10T00:00:00+00:00",
    )

    def type_metric(metric_name: str, value: float, observations: int, denominator: str) -> CorpusMetricRecord:
        return CorpusMetricRecord(
            run_id="run-type",
            text_id=text.text_id,
            text_version_id=text.text_version_id,
            title=text.title,
            author="",
            collection="",
            date_label="",
            genre="",
            lexicon_id="fixture-vad",
            lexicon="Fixture VAD",
            value_kind="vad",
            metric=metric_name,
            dimension="valence" if metric_name == "vad_mean" else "",
            category="",
            weighting="type",
            scale="normalized_0_1",
            denominator=denominator,
            value=value,
            observations=observations,
            matched_tokens=8,
            lexical_tokens=22,
            coverage=8 / 22,
            completed_at="2026-08-10T00:00:00+00:00",
            analysis_view="content_words",
        )

    archive_bytes = build_corpus_export_bundle(
        project,
        (text,),
        (
            type_metric("type_coverage", 7 / 21, 7, "21 eligible types"),
            type_metric("vad_mean", 0.5, 7, "7 observations; 8/22 eligible tokens matched"),
        ),
        (),
    )
    with zipfile.ZipFile(io.BytesIO(archive_bytes)) as archive:
        rows = list(
            csv.DictReader(
                io.StringIO(
                    archive.read("04_AUDIT/corpus_vad_metrics.csv").decode("utf-8-sig")
                )
            )
        )
        mean_row = next(row for row in rows if row["metric"] == "vad_mean")
        assert mean_row["denominator"] == "7 observations; 7/21 eligible types matched"
        assert mean_row["matched_tokens"] == "7"
        assert mean_row["lexical_tokens"] == "21"
        assert float(mean_row["coverage"]) == pytest.approx(7 / 21)

        report = Document(io.BytesIO(archive.read("01_REPORTS/Corpus_Report.docx")))
        report_text = "\n".join(
            cell.text
            for table in report.tables
            for row in table.rows
            for cell in row.cells
        )
        assert "Matched" in report_text
        assert "7" in report_text


def test_corpus_current_view_uses_content_scope_only_for_overridden_module() -> None:
    project = ProjectRecord(
        project_id="project-override",
        title="Override Project",
        description="",
        researcher="Researcher",
        created_at="2026-08-11T00:00:00+00:00",
        updated_at="2026-08-11T00:00:00+00:00",
    )
    text = CorpusTextRecord(
        text_id="override-text",
        text_version_id="override-text:v1",
        project_id=project.project_id,
        title="Fixture",
        source_name="fixture.txt",
        relative_path="fixture.txt",
        author="",
        collection="",
        date_label="",
        genre="",
        notes="",
        custom_metadata={},
        original_text="Fixture.",
        text_sha256="d" * 64,
        imported_at="2026-08-11T00:00:00+00:00",
        updated_at="2026-08-11T00:00:00+00:00",
    )

    def metric(name: str, view: str, value: float) -> CorpusMetricRecord:
        return CorpusMetricRecord(
            run_id="run-override",
            text_id=text.text_id,
            text_version_id=text.text_version_id,
            title=text.title,
            author="",
            collection="",
            date_label="",
            genre="",
            lexicon_id="fixture",
            lexicon="Fixture Resource",
            value_kind="continuous",
            metric=name,
            dimension="",
            category="",
            weighting="token",
            scale="source",
            denominator="1 observation",
            value=value,
            observations=1,
            matched_tokens=1,
            lexical_tokens=1,
            coverage=1.0,
            completed_at="2026-08-11T00:00:00+00:00",
            analysis_view=view,
        )

    archive_bytes = build_corpus_export_bundle(
        project,
        (text,),
        (
            metric("frequency_frequency_mean", "stopwords_excluded", 4.0),
            metric("frequency_frequency_mean", "content_words", 5.0),
            metric("concreteness_concreteness_mean", "stopwords_excluded", 2.0),
            metric("concreteness_concreteness_mean", "content_words", 3.0),
        ),
        (),
        profile_selection=ProfileSelection(),
        export_mode="current_view",
        report_section="Lexical Character, Imagery & Embodiment",
        module_scope_overrides=frozenset({"frequency"}),
    )

    with zipfile.ZipFile(io.BytesIO(archive_bytes)) as archive:
        rows = list(
            csv.DictReader(
                io.StringIO(
                    archive.read("04_AUDIT/corpus_vad_metrics.csv").decode("utf-8-sig")
                )
            )
        )
        assert {(row["metric"], row["analysis_view"]) for row in rows} == {
            ("frequency_frequency_mean", "content_words"),
            ("concreteness_concreteness_mean", "stopwords_excluded"),
        }
        assert "05_REPRODUCIBILITY/Module_Scope_Overrides.csv" in archive.namelist()
        report = Document(io.BytesIO(archive.read("01_REPORTS/Corpus_Report.docx")))
        report_text = "\n".join(
            [
                *(paragraph.text for paragraph in report.paragraphs),
                *(
                    cell.text
                    for table in report.tables
                    for table_row in table.rows
                    for cell in table_row.cells
                ),
            ]
        )
        assert "Mean Zipf Frequency" in report_text
        assert "Mean Concreteness" in report_text
        assert "Frequency Frequency Mean" not in report_text
        assert "Concreteness Concreteness Mean" not in report_text
