from __future__ import annotations

import io

from docx import Document

from versevad.exports.docx_report import build_comprehensive_analysis_report


def _csv(text: str) -> bytes:
    return text.strip().encode("utf-8") + b"\n"


def _document_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join((*paragraphs, *cells))


def _profile_csv() -> bytes:
    return _csv(
        """
profile_id,scope,weighting,module_id,source_id,source,metric_id,metric,value,median,population_standard_deviation,first_quartile,third_quartile,minimum,maximum,cumulative_value,value_per_100_observations,above_midpoint_load,below_midpoint_load,net_midpoint_load,absolute_midpoint_load,average_deviation_from_mean,observation_count,eligible_token_count,matched_token_count,unmatched_token_count,token_coverage,eligible_type_count,matched_type_count,unmatched_type_count,type_coverage,excluded_stopword_count,excluded_non_content_count,phrase_match_count,type_identity_rule,unit
stopword_excluded-token_weighted,stopword_excluded,token_weighted,vad,nrc_vad_v2_1,NRC VAD Lexicon v2.1,valence,Mean normative valence,0.572345,0.646111,0.188812,0.421111,0.711111,0.111111,0.922222,2.345678,8.687696,1.234567,0.567891,0.666676,1.802458,0.145678,27,31,27,4,0.870967,24,22,2,0.916667,5,0,1,lemma,normalized 0-1
stopword_excluded-token_weighted,stopword_excluded,token_weighted,concreteness,brysbaert_concreteness,Brysbaert ratings,concreteness,Mean concreteness,3.255555,3.100001,0.912345,2.500001,3.800001,1.500001,4.900001,97.666650,325.555500,,,,,0.711111,30,31,30,1,0.967742,24,23,1,0.958333,5,0,0,lemma,source 1-5
"""
    )


def test_current_view_comprehensive_report_is_readable_and_deterministic() -> None:
    files = {
        "profile_metrics_selected.csv": _profile_csv(),
        "readability_summary.csv": _csv(
            "section,metric,value,unit_or_scale,denominator,note\n"
            "VerseVAD Poetic Reading Ease,VV-PRE score,76.42891,0-100,one completed text,Accessible\n"
        ),
        "phase2_match_audit.csv": _csv("token,value\nhope,0.9\n"),
    }
    kwargs = dict(
        export_files=files,
        text_title="Example Poem",
        author="Example Author",
        analysis_timestamp="2026-08-09T08:30:00-04:00",
        export_mode="current_view",
        visible_section="Affective Evidence",
        workspace_label="Single Poem",
        text_id="text-1",
        result_id="version-1",
        source_sha256="abc123",
        analysis_profiles=("stopword_excluded-token_weighted",),
        active_preset="Full Poetic Analysis",
        software_version="1.0.0",
    )
    first = build_comprehensive_analysis_report(**kwargs)
    second = build_comprehensive_analysis_report(**kwargs)
    assert first == second
    text = _document_text(first)
    assert "Computational Poetics\nAnalysis Report" in text
    assert "Current View Report" in text
    assert "Example Author" in text
    assert "[Enter analyst name]" in text
    assert "[Enter research question]" in text
    assert "0.572" in text
    assert "Valence, Arousal, and Dominance" in text
    assert "Not calculated" in text


def test_complete_audit_report_marks_uncalculated_modules_and_lists_audit_files() -> None:
    content = build_comprehensive_analysis_report(
        export_files={
            "profile_metrics_selected.csv": _profile_csv(),
            "profile_metrics_all_compatible.csv": _profile_csv(),
            "phase2_token_audit.csv": _csv("token,value\nhope,0.9\n"),
        },
        text_title="Example Poem",
        export_mode="complete_audit",
        software_version="1.0.0",
    )
    text = _document_text(content)
    assert "Complete Audit Report" in text
    assert "All Compatible Lexical Profiles" in text
    assert "Not calculated" in text
    assert "phase2_token_audit.csv" in text
    assert "full-precision" in text.lower()


def test_current_view_distinguishes_calculated_but_unreported_modules() -> None:
    content = build_comprehensive_analysis_report(
        export_files={"profile_metrics_selected.csv": _profile_csv()},
        text_title="Example Poem",
        export_mode="current_view",
        visible_section="Overview",
        calculated_modules=("vad", "readability", "meter", "poetry_id"),
    )
    text = _document_text(content)
    assert "Calculated, not included" in text
    assert "selected Overview report section" in text


def test_corpus_profile_appendix_names_real_standardized_companion_file() -> None:
    content = build_comprehensive_analysis_report(
        export_files={
            "profile_metrics_selected.csv": _profile_csv(),
            "profile_metrics_all_compatible.csv": _profile_csv(),
            "corpus_vad_profiles.csv": _csv("dimension,value\nvalence,0.5\n"),
        },
        text_title="Example Corpus",
        export_mode="complete_audit",
        workspace_label="Saved Projects / Corpus",
    )
    text = _document_text(content)
    assert "All Compatible Lexical Profiles" in text
    assert "03_MASTER_DATA/All_Profiles.csv" in text
    assert "companion corpus_vad_profiles.csv" not in text


def test_primary_profile_table_uses_weighting_appropriate_coverage() -> None:
    type_profile = _profile_csv().decode("utf-8").replace(
        "stopword_excluded-token_weighted,stopword_excluded,token_weighted",
        "stopword_excluded-type_weighted,stopword_excluded,type_weighted",
    ).encode("utf-8")
    content = build_comprehensive_analysis_report(
        export_files={"profile_metrics_selected.csv": type_profile},
        text_title="Example Poem",
        export_mode="current_view",
    )
    document = Document(io.BytesIO(content))
    primary_table = next(
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text == "Source"
        and any(cell.text == "Coverage" for cell in table.rows[0].cells)
    )
    headings = [cell.text for cell in primary_table.rows[0].cells]
    coverage_index = headings.index("Coverage")
    assert primary_table.rows[1].cells[coverage_index].text == "91.7%"


def test_corpus_report_distinguishes_pooled_and_between_work_dispersion() -> None:
    content = build_comprehensive_analysis_report(
        export_files={"profile_metrics_selected.csv": _profile_csv()},
        text_title="Example Corpus",
        export_mode="complete_audit",
        workspace_label="Saved Projects / Corpus",
    )
    text = _document_text(content)
    assert "Lexical and Between-Work Dispersion" in text
    assert "pooled matched lexical ratings" in text
    assert "between-work dispersion" in text


def test_comprehensive_report_includes_experiential_dynamics_when_completed() -> None:
    summary = _csv(
        "assessment_id,text_version_id,assessment_timing,submitted_at,methodology_version,questionnaire_version,configuration_id,fixed_scope,fixed_weighting,agreement_tolerance,dynamic_signature,compact_code,dimension,dimension_label,resource_id,resource,resource_version,resource_sha256,measured_source_value,measured_source_unit,measured_normalized_0_1,experienced_raw_mean_1_5,experienced_normalized_0_1,reader_response_population_sd,dynamic_gap_experienced_minus_measured,relationship,relationship_label,compact_symbol,eligible_token_count,matched_token_count,token_coverage\n"
        "assessment-1,text-1,pre_analysis,2026-08-13T12:00:00+00:00,experiential-dynamics-1.0,reader-response-16-v1,config-1,stopword_excluded,token_weighted,0.1,Darkened · Charged · Constrained · Evoked,V↓ A↑ D↓ C↑,valence,Valence,nrc_vad_v2_1,NRC VAD Lexicon v2.1,2.1,abc,0.4,normalized 0-1,0.4,1.8,0.2,0.1,-0.2,experienced_lower,Darkened,↓,20,18,0.9\n"
    )
    responses = _csv(
        "assessment_id,text_version_id,assessment_timing,item_id,dimension,prompt,numeric_response_1_5,selected_response\n"
        "assessment-1,text-1,pre_analysis,V1,valence,Question,2,Somewhat negative\n"
    )
    content = build_comprehensive_analysis_report(
        export_files={
            "profile_metrics_selected.csv": _profile_csv(),
            "experiential_dynamics_summary.csv": summary,
            "experiential_dynamics_responses.csv": responses,
        },
        text_title="Example Poem",
        export_mode="current_view",
        visible_section="Affective Evidence",
        calculated_modules=("vad", "experiential_dynamics"),
    )
    text = _document_text(content)
    assert "Experiential Dynamics" in text
    assert "Darkened · Charged · Constrained · Evoked" in text
    assert "Reader-derived responses" in text
    assert "The reader-derived assessment and fixed lexical comparison" in text
